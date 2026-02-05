"""
نافذة استمارة إتلاف الوثائق
Destruction Form Dialog

نافذة لإنشاء وطباعة وتصدير استمارة إتلاف الوثائق
"""

import os

from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QPushButton, QTableWidget, QTableWidgetItem, QGroupBox,
    QFormLayout, QFileDialog, QMessageBox
)
from PyQt6.QtCore import Qt, QRectF, QMarginsF
from PyQt6.QtGui import QPainter, QFont, QColor, QPen, QImage, QPageSize, QPageLayout

from ..constants import COLORS, FONT_SIZES, DIMENSIONS, ICONS
from ..ui_styles import TITLE_STYLES, BUTTON_STYLES, PAGES_INFO_STYLE


class DestructionFormDialog(QDialog):
    """
    نافذة استمارة إتلاف الوثائق
    
    تتيح هذه النافذة:
    - إدخال معلومات الجهة
    - تحميل الوثائق المحددة
    - طباعة الاستمارة
    - تصدير إلى Excel و Word
    """
    
    ROWS_PER_PAGE = DIMENSIONS.ROWS_PER_PAGE  # عدد الصفوف في كل صفحة
    
    def __init__(self, parent=None, db=None, selected_docs=None):
        """
        تهيئة النافذة
        
        Args:
            parent: النافذة الأب
            db: مدير قاعدة البيانات
            selected_docs: الوثائق المحددة لإضافتها للاستمارة
        """
        super().__init__(parent)
        self.db = db
        self.selected_docs = selected_docs or []
        
        self.setWindowTitle('استمارة إتلاف الوثائق')
        self.setMinimumSize(900, 700)
        
        self._init_ui()
        self._load_selected_documents()
        self._update_pages_info()
    
    def _init_ui(self):
        """إنشاء واجهة المستخدم"""
        layout = QVBoxLayout()
        
        # عنوان النموذج
        self._create_title(layout)
        
        # معلومات الرأس
        self._create_header_fields(layout)
        
        # معلومات الصفحات
        self._create_pages_info(layout)
        
        # جدول الوثائق
        self._create_documents_table(layout)
        
        # أزرار الإجراءات
        self._create_action_buttons(layout)
        
        self.setLayout(layout)
    
    def _create_title(self, parent_layout):
        """إنشاء عنوان النموذج"""
        title = QLabel(f'{ICONS.DOCUMENT} استمارة إتلاف الوثائق')
        title.setStyleSheet(TITLE_STYLES['large'])
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        parent_layout.addWidget(title)
    
    def _create_header_fields(self, parent_layout):
        """إنشاء حقول معلومات الرأس"""
        header_group = QGroupBox('معلومات الجهة')
        header_layout = QFormLayout()
        
        self.agency_input = QLineEdit()
        self.agency_input.setPlaceholderText('أدخل اسم الوكالة')
        header_layout.addRow('الوكالة:', self.agency_input)
        
        self.directorate_input = QLineEdit()
        self.directorate_input.setPlaceholderText('أدخل اسم التشكيل أو المديرية')
        header_layout.addRow('التشكيل/المديرية:', self.directorate_input)
        
        self.section_input = QLineEdit()
        self.section_input.setPlaceholderText('أدخل اسم القسم')
        header_layout.addRow('القسم:', self.section_input)
        
        self.division_input = QLineEdit()
        self.division_input.setPlaceholderText('أدخل اسم الشعبة')
        header_layout.addRow('الشعبة:', self.division_input)
        
        header_group.setLayout(header_layout)
        parent_layout.addWidget(header_group)
    
    def _create_pages_info(self, parent_layout):
        """إنشاء معلومات الصفحات"""
        pages_info_layout = QHBoxLayout()
        
        self.pages_info_label = QLabel('')
        self.pages_info_label.setStyleSheet(PAGES_INFO_STYLE)
        pages_info_layout.addWidget(self.pages_info_label)
        pages_info_layout.addStretch()
        
        parent_layout.addLayout(pages_info_layout)
    
    def _create_documents_table(self, parent_layout):
        """إنشاء جدول الوثائق"""
        docs_group = QGroupBox('الوثائق المراد إتلافها')
        docs_layout = QVBoxLayout()
        
        # أزرار التحكم بالجدول
        table_buttons = QHBoxLayout()
        
        add_row_btn = QPushButton(f'{ICONS.ADD} إضافة صف')
        add_row_btn.clicked.connect(self._add_row)
        table_buttons.addWidget(add_row_btn)
        
        remove_row_btn = QPushButton(f'➖ حذف صف')
        remove_row_btn.clicked.connect(self._remove_row)
        table_buttons.addWidget(remove_row_btn)
        
        table_buttons.addStretch()
        
        load_selected_btn = QPushButton(f'{ICONS.IMPORT} تحميل الوثائق المحددة')
        load_selected_btn.clicked.connect(self._load_selected_documents)
        table_buttons.addWidget(load_selected_btn)
        
        docs_layout.addLayout(table_buttons)
        
        # الجدول
        self.docs_table = QTableWidget()
        self.docs_table.setColumnCount(7)
        self.docs_table.setHorizontalHeaderLabels([
            'ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 
            'مضمونها', 'تصنيف الوثيقة\n(أ، ب، ج)', 'الفقرة القانونية'
        ])
        
        # تعيين عرض الأعمدة
        self.docs_table.setColumnWidth(0, 40)
        self.docs_table.setColumnWidth(1, 100)
        self.docs_table.setColumnWidth(2, 100)
        self.docs_table.setColumnWidth(3, 120)
        self.docs_table.setColumnWidth(4, 200)
        self.docs_table.setColumnWidth(5, 80)
        self.docs_table.setColumnWidth(6, 150)
        
        self.docs_table.setAlternatingRowColors(True)
        self.docs_table.model().rowsInserted.connect(self._update_pages_info)
        self.docs_table.model().rowsRemoved.connect(self._update_pages_info)
        
        docs_layout.addWidget(self.docs_table)
        docs_group.setLayout(docs_layout)
        parent_layout.addWidget(docs_group)
    
    def _create_action_buttons(self, parent_layout):
        """إنشاء أزرار الإجراءات"""
        buttons_layout = QHBoxLayout()
        
        # زر الطباعة
        print_btn = QPushButton(f'{ICONS.PRINT} طباعة')
        print_btn.setStyleSheet(BUTTON_STYLES['purple'])
        print_btn.clicked.connect(self._print_form)
        buttons_layout.addWidget(print_btn)
        
        # زر تصدير Excel
        export_excel_btn = QPushButton(f'{ICONS.EXCEL} تصدير Excel')
        export_excel_btn.setStyleSheet(BUTTON_STYLES['success'])
        export_excel_btn.clicked.connect(self._export_to_excel)
        buttons_layout.addWidget(export_excel_btn)
        
        # زر تصدير Word
        export_word_btn = QPushButton(f'{ICONS.WORD} تصدير Word')
        export_word_btn.setStyleSheet(BUTTON_STYLES['info'])
        export_word_btn.clicked.connect(self._export_to_word)
        buttons_layout.addWidget(export_word_btn)
        
        buttons_layout.addStretch()
        
        # زر الإغلاق
        close_btn = QPushButton(f'{ICONS.CANCEL} إغلاق')
        close_btn.clicked.connect(self.reject)
        buttons_layout.addWidget(close_btn)
        
        parent_layout.addLayout(buttons_layout)
    
    # =========================================================================
    # وظائف إدارة الجدول
    # =========================================================================
    
    def _update_pages_info(self):
        """تحديث معلومات الصفحات"""
        total_docs = self.docs_table.rowCount()
        total_pages = (total_docs + self.ROWS_PER_PAGE - 1) // self.ROWS_PER_PAGE if total_docs > 0 else 1
        self.pages_info_label.setText(
            f'{ICONS.FILE} إجمالي الوثائق: {total_docs} | '
            f'عدد الصفحات: {total_pages} (25 وثيقة لكل صفحة)'
        )
    
    def _add_row(self):
        """إضافة صف جديد"""
        row = self.docs_table.rowCount()
        self.docs_table.insertRow(row)
        self.docs_table.setItem(row, 0, QTableWidgetItem(str(row + 1)))
    
    def _remove_row(self):
        """حذف الصف المحدد"""
        current_row = self.docs_table.currentRow()
        if current_row >= 0:
            self.docs_table.removeRow(current_row)
            # تحديث أرقام التسلسل
            for i in range(self.docs_table.rowCount()):
                self.docs_table.setItem(i, 0, QTableWidgetItem(str(i + 1)))
    
    def _load_selected_documents(self):
        """تحميل الوثائق المحددة من النافذة الرئيسية"""
        if not self.selected_docs:
            return
        
        self.docs_table.setRowCount(0)
        
        for idx, doc in enumerate(self.selected_docs):
            row = self.docs_table.rowCount()
            self.docs_table.insertRow(row)
            
            # ت (التسلسل)
            self.docs_table.setItem(row, 0, QTableWidgetItem(str(idx + 1)))
            
            # رقم الوثيقة
            doc_name = doc[1] or ''
            doc_number = doc_name.split()[0] if doc_name else ''
            self.docs_table.setItem(row, 1, QTableWidgetItem(doc_number))
            
            # تاريخها
            self.docs_table.setItem(row, 2, QTableWidgetItem(doc[2] or ''))
            
            # جهة الإصدار
            self.docs_table.setItem(row, 3, QTableWidgetItem(doc[4] or ''))
            
            # مضمونها
            self.docs_table.setItem(row, 4, QTableWidgetItem(doc[3] or ''))
            
            # تصنيف الوثيقة
            self.docs_table.setItem(row, 5, QTableWidgetItem(doc[5] or ''))
            
            # الفقرة القانونية
            self.docs_table.setItem(row, 6, QTableWidgetItem(doc[6] or ''))
        
        self._update_pages_info()
    
    def _get_table_data(self):
        """جمع بيانات الجدول"""
        data = []
        for row in range(self.docs_table.rowCount()):
            row_data = []
            for col in range(self.docs_table.columnCount()):
                item = self.docs_table.item(row, col)
                row_data.append(item.text() if item else '')
            data.append(row_data)
        return data
    
    def _split_data_into_pages(self, data):
        """تقسيم البيانات إلى صفحات (25 صف لكل صفحة)"""
        pages = []
        for i in range(0, len(data), self.ROWS_PER_PAGE):
            page_data = data[i:i + self.ROWS_PER_PAGE]
            # إضافة صفوف فارغة إذا كانت الصفحة غير مكتملة
            while len(page_data) < self.ROWS_PER_PAGE:
                empty_row = [str(len(page_data) + 1 + i)] + [''] * 6
                page_data.append(empty_row)
            pages.append(page_data)
        
        # إذا لم تكن هناك بيانات، أنشئ صفحة فارغة
        if not pages:
            page_data = []
            for j in range(self.ROWS_PER_PAGE):
                page_data.append([str(j + 1)] + [''] * 6)
            pages.append(page_data)
        
        return pages
    
    # =========================================================================
    # وظائف الطباعة
    # =========================================================================
    
    def _print_form(self):
        """طباعة الاستمارة مع معاينة باستخدام QPainter"""
        try:
            from PyQt6.QtPrintSupport import QPrinter, QPrintPreviewDialog
            
            # جمع البيانات وتقسيمها
            table_data = self._get_table_data()
            self._print_pages = self._split_data_into_pages(table_data)
            self._total_pages = len(self._print_pages)
            
            # حفظ معلومات النموذج
            self._form_info = {
                'agency': self.agency_input.text(),
                'directorate': self.directorate_input.text(),
                'section': self.section_input.text(),
                'division': self.division_input.text()
            }
            
            # إعداد الطابعة
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            page_size = QPageSize(QPageSize.PageSizeId.A4)
            page_layout = QPageLayout(
                page_size, 
                QPageLayout.Orientation.Portrait, 
                QMarginsF(15, 15, 15, 15)
            )
            printer.setPageLayout(page_layout)
            
            # إنشاء نافذة المعاينة
            preview = QPrintPreviewDialog(printer, self)
            preview.setWindowTitle(f'معاينة الطباعة - {self._total_pages} صفحة')
            preview.setMinimumSize(800, 600)
            
            preview.paintRequested.connect(self._draw_pages_with_painter)
            preview.exec()
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء الطباعة:\n{str(e)}')
    
    def _draw_pages_with_painter(self, printer):
        """رسم الصفحات باستخدام QPainter للتحكم الكامل"""
        from PyQt6.QtPrintSupport import QPrinter
        
        painter = QPainter()
        painter.begin(printer)
        
        # استخدام وحدة Millimeter للحصول على أبعاد حقيقية
        page_rect = printer.pageRect(QPrinter.Unit.Millimeter)
        width_mm = page_rect.width()
        height_mm = page_rect.height()
        
        # تحويل من مليمتر إلى بكسل
        dpi = printer.resolution()
        px_per_mm = dpi / 25.4
        
        width = width_mm * px_per_mm
        height = height_mm * px_per_mm
        
        # هوامش بالمليمتر ثم تحويلها
        margin_mm = 10
        margin = margin_mm * px_per_mm
        content_width = width - (2 * margin)
        
        # ارتفاع الصف (حوالي 7mm للصف)
        row_height = 7 * px_per_mm
        
        for page_idx, page_data in enumerate(self._print_pages):
            if page_idx > 0:
                printer.newPage()
            
            y_pos = margin
            
            # ===== الشعار (أعلى وسط) =====
            logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "..", "MOI.png")
            if os.path.exists(logo_path):
                logo = QImage(logo_path)
                if not logo.isNull():
                    logo_size = 25 * px_per_mm
                    logo_x = margin + (content_width - logo_size) / 2
                    logo_y = y_pos
                    painter.drawImage(QRectF(logo_x, logo_y, logo_size, logo_size), logo)
            
            # ===== معلومات الوزارة (أعلى يمين) =====
            ministry_font = QFont("Arial")
            ministry_font.setPointSize(10)
            ministry_font.setBold(True)
            painter.setFont(ministry_font)
            painter.setPen(Qt.GlobalColor.black)
            
            ministry_lines = [
                "وزارة الداخلية",
                "وكالة الوزارة لشؤون الإدارية والمالية",
                "مديرية إدارة الموارد البشرية",
                "مديرية السجلات والوثائق"
            ]
            
            line_h = 5 * px_per_mm
            temp_y = y_pos
            for line in ministry_lines:
                ministry_rect = QRectF(margin + content_width * 0.55, temp_y, content_width * 0.45, line_h)
                painter.drawText(ministry_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, line)
                temp_y += line_h
            
            # ===== معلومات الإصدار (أعلى يسار) =====
            version_font = QFont("Arial")
            version_font.setPointSize(8)
            painter.setFont(version_font)
            
            version_data = [
                ("رقم الإصدار", "0.1"),
                ("سنة الإصدار", "2023"),
                ("رقم الترميز", "م.ب-س"),
                ("نموذج", "(37)")
            ]
            
            inner_y = y_pos
            for label, value in version_data:
                val_rect = QRectF(margin, inner_y, 15 * px_per_mm, line_h)
                painter.drawText(val_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, value)
                slash_rect = QRectF(margin + 15 * px_per_mm, inner_y, 3 * px_per_mm, line_h)
                painter.drawText(slash_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, "/")
                lbl_rect = QRectF(margin + 18 * px_per_mm, inner_y, 25 * px_per_mm, line_h)
                painter.drawText(lbl_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, label)
                inner_y += line_h
            
            y_pos = temp_y + 5 * px_per_mm
            
            # ===== العنوان الرئيسي (وسط) =====
            title_font = QFont("Arial")
            title_font.setPointSize(16)
            title_font.setBold(True)
            painter.setFont(title_font)
            painter.setPen(Qt.GlobalColor.black)
            
            title_height = 10 * px_per_mm
            title_rect = QRectF(margin, y_pos, content_width, title_height)
            painter.drawText(title_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, "استمارة إتلاف الوثائق")
            y_pos += title_height + (4 * px_per_mm)
            
            # ===== معلومات النموذج =====
            info_font = QFont("Arial")
            info_font.setPointSize(10)
            painter.setFont(info_font)
            
            info_data = [
                ("الوكالة", self._form_info['agency']),
                ("التشكيل أو المديرية", self._form_info['directorate']),
                ("القسم", self._form_info['section']),
                ("الشعبة", self._form_info['division'])
            ]
            
            line_height = 5 * px_per_mm
            label_w = 32 * px_per_mm
            
            for label, value in info_data:
                lbl_rect = QRectF(margin + content_width - label_w, y_pos, label_w, line_height)
                painter.drawText(lbl_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, label)
                colon_rect = QRectF(margin + content_width - label_w - (4 * px_per_mm), y_pos, 4 * px_per_mm, line_height)
                painter.drawText(colon_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, ":")
                val_rect = QRectF(margin, y_pos, content_width - label_w - (6 * px_per_mm), line_height)
                painter.drawText(val_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, value)
                y_pos += line_height
            
            y_pos += 3 * px_per_mm
            
            # ===== الجدول =====
            col_widths = [0.04, 0.08, 0.10, 0.12, 0.34, 0.08, 0.24]
            headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف', 'الفقرة القانونية']
            
            # رسم رأس الجدول
            header_font = QFont("Arial")
            header_font.setPointSize(8)
            header_font.setBold(True)
            painter.setFont(header_font)
            
            painter.setBrush(QColor(217, 225, 242))
            painter.setPen(QPen(Qt.GlobalColor.black, 1))
            painter.drawRect(QRectF(margin, y_pos, content_width, row_height))
            
            x_pos = margin + content_width
            for header, col_w in zip(headers, col_widths):
                cell_width = content_width * col_w
                x_pos -= cell_width
                cell_rect = QRectF(x_pos, y_pos, cell_width, row_height)
                painter.drawRect(cell_rect)
                painter.drawText(cell_rect, Qt.AlignmentFlag.AlignCenter, header)
            
            y_pos += row_height
            
            # رسم صفوف البيانات
            data_font = QFont("Arial")
            data_font.setPointSize(7)
            painter.setFont(data_font)
            painter.setBrush(Qt.GlobalColor.white)
            
            for row_idx, row in enumerate(page_data, 1):
                x_pos = margin + content_width
                row_data = [str(row_idx)] + list(row[1:])
                
                for cell, col_w in zip(row_data, col_widths):
                    cell_width = content_width * col_w
                    x_pos -= cell_width
                    cell_rect = QRectF(x_pos, y_pos, cell_width, row_height)
                    painter.drawRect(cell_rect)
                    display_text = str(cell)[:45] if len(str(cell)) > 45 else str(cell)
                    painter.drawText(cell_rect.adjusted(2, 0, -2, 0), Qt.AlignmentFlag.AlignCenter, display_text)
                
                y_pos += row_height
            
            # ===== رقم الصفحة =====
            page_font = QFont("Arial")
            page_font.setPointSize(8)
            painter.setFont(page_font)
            page_num_text = f"صفحة {page_idx + 1} من {self._total_pages}"
            page_rect_bottom = QRectF(margin, height - margin - (5 * px_per_mm), content_width, 5 * px_per_mm)
            painter.drawText(page_rect_bottom, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, page_num_text)
        
        painter.end()
    
    # =========================================================================
    # وظائف التصدير
    # =========================================================================
    
    def _export_to_excel(self):
        """تصدير إلى ملف Excel"""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        except ImportError:
            QMessageBox.warning(
                self, 'خطأ', 
                'مكتبة openpyxl غير مثبتة!\nقم بتثبيتها: pip install openpyxl'
            )
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'حفظ استمارة الإتلاف',
            'استمارة_اتلاف_الوثائق.xlsx',
            'Excel Files (*.xlsx)'
        )
        
        if not file_path:
            return
        
        try:
            table_data = self._get_table_data()
            pages = self._split_data_into_pages(table_data)
            
            wb = openpyxl.Workbook()
            wb.remove(wb.active)
            
            title_font = Font(size=16, bold=True)
            header_font = Font(size=11, bold=True)
            cell_font = Font(size=10)
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            right_align = Alignment(horizontal='right', vertical='center')
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
            
            for page_num, page_data in enumerate(pages, 1):
                ws = wb.create_sheet(title=f'صفحة {page_num}')
                ws.sheet_view.rightToLeft = True
                
                ws.merge_cells('A1:G1')
                ws['A1'] = 'استمارة إتلاف الوثائق'
                ws['A1'].font = title_font
                ws['A1'].alignment = center_align
                ws.row_dimensions[1].height = 25
                
                ws['A3'] = f'الوكالة: {self.agency_input.text()}'
                ws['A3'].alignment = right_align
                ws.merge_cells('A3:G3')
                
                ws['A4'] = f'التشكيل أو المديرية: {self.directorate_input.text()}'
                ws['A4'].alignment = right_align
                ws.merge_cells('A4:G4')
                
                ws['A5'] = f'القسم: {self.section_input.text()}'
                ws['A5'].alignment = right_align
                ws.merge_cells('A5:G5')
                
                ws['A6'] = f'الشعبة: {self.division_input.text()}'
                ws['A6'].alignment = right_align
                ws.merge_cells('A6:G6')
                
                ws['G7'] = f'صفحة {page_num} من {len(pages)}'
                ws['G7'].alignment = Alignment(horizontal='left', vertical='center')
                
                headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف\n(أ،ب،ج)', 'الفقرة القانونية']
                for col_idx, header in enumerate(headers, 1):
                    cell = ws.cell(row=9, column=col_idx, value=header)
                    cell.font = header_font
                    cell.alignment = center_align
                    cell.border = thin_border
                    cell.fill = header_fill
                
                ws.column_dimensions['A'].width = 5
                ws.column_dimensions['B'].width = 12
                ws.column_dimensions['C'].width = 12
                ws.column_dimensions['D'].width = 18
                ws.column_dimensions['E'].width = 30
                ws.column_dimensions['F'].width = 10
                ws.column_dimensions['G'].width = 25
                
                for row_idx, row_data in enumerate(page_data, 10):
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = cell_font
                        cell.alignment = center_align
                        cell.border = thin_border
            
            wb.save(file_path)
            QMessageBox.information(self, 'نجح', f'تم تصدير {len(pages)} صفحة إلى:\n{file_path}')
            
            # فتح الملف (Windows فقط)
            try:
                os.startfile(file_path)
            except:
                pass
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء التصدير:\n{str(e)}')
    
    def _export_to_word(self):
        """تصدير إلى ملف Word"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            QMessageBox.warning(
                self, 'خطأ', 
                'مكتبة python-docx غير مثبتة!\nقم بتثبيتها: pip install python-docx'
            )
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'حفظ استمارة الإتلاف',
            'استمارة_اتلاف_الوثائق.docx',
            'Word Files (*.docx)'
        )
        
        if not file_path:
            return
        
        try:
            table_data = self._get_table_data()
            pages = self._split_data_into_pages(table_data)
            
            doc = Document()
            
            for page_num, page_data in enumerate(pages, 1):
                if page_num == 1:
                    section = doc.sections[0]
                else:
                    section = doc.add_section()
                
                sectPr = section._sectPr
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                sectPr.append(bidi)
                
                title = doc.add_paragraph('استمارة إتلاف الوثائق')
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].bold = True
                title.runs[0].font.size = Pt(16)
                
                doc.add_paragraph(f'الوكالة: {self.agency_input.text()}')
                doc.add_paragraph(f'التشكيل أو المديرية: {self.directorate_input.text()}')
                doc.add_paragraph(f'القسم: {self.section_input.text()}')
                doc.add_paragraph(f'الشعبة: {self.division_input.text()}')
                
                page_info = doc.add_paragraph(f'صفحة {page_num} من {len(pages)}')
                page_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                table = doc.add_table(rows=len(page_data) + 1, cols=7)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف\n(أ،ب،ج)', 'الفقرة القانونية']
                header_row = table.rows[0]
                for idx, header in enumerate(headers):
                    cell = header_row.cells[idx]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                
                for row_idx, row_data in enumerate(page_data):
                    row = table.rows[row_idx + 1]
                    for col_idx, value in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = value
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)
                
                if page_num < len(pages):
                    doc.add_page_break()
            
            doc.save(file_path)
            QMessageBox.information(self, 'نجح', f'تم تصدير {len(pages)} صفحة إلى:\n{file_path}')
            
            # فتح الملف (Windows فقط)
            try:
                os.startfile(file_path)
            except:
                pass
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء التصدير:\n{str(e)}')
