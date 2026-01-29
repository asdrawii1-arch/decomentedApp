import sys
import os
import tempfile
from pathlib import Path
from datetime import datetime
from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QTableWidget, 
    QTableWidgetItem, QPushButton, QLineEdit, QLabel, QFileDialog,
    QDialog, QDialogButtonBox, QComboBox, QSpinBox, QMessageBox,
    QTabWidget, QGroupBox, QFormLayout, QTextEdit, QListWidget,
    QListWidgetItem, QProgressBar, QProgressDialog, QCheckBox
)
from PyQt6.QtCore import Qt, QSize, pyqtSignal, QTimer
from PyQt6.QtGui import QIcon, QFont, QColor
from PyQt6.QtWidgets import QApplication

# إضافة المسارات
sys.path.insert(0, str(Path(__file__).parent / 'src'))

# التحقق من توفر مكتبة السكانر
SCANNER_AVAILABLE = False
SCANNER_COUNT = 0
try:
    import win32com.client
    SCANNER_AVAILABLE = True
    # التحقق من عدد السكانرات المتصلة
    try:
        _wia_manager = win32com.client.Dispatch("WIA.DeviceManager")
        SCANNER_COUNT = _wia_manager.DeviceInfos.Count
    except:
        SCANNER_COUNT = 0
except ImportError:
    SCANNER_AVAILABLE = False
    SCANNER_COUNT = 0

from database.db_manager import DatabaseManager
from app.filename_parser import FilenameParser, ImageSequenceHandler
from app.ui_styles import MAIN_STYLESHEET, COLORS, SIZES
from app.image_manager import ImageManager
from app.document_viewer import DocumentViewerWindow
from app.helpers import ValidationHelper, DateHelper, ExportHelper, DatabaseBackupHelper

# استيراد OCR اختياري
try:
    from app.ocr_extractor import OCRExtractor
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("[WARNING] مكتبة easyocr غير مثبتة - ميزة استخراج المضمون غير متاحة")


def choose_year_folder(parent):
    """دالة مساعدة لعرض مجلدات السنوات الموجودة أو إنشاء مجلد سنة جديدة.
    ترجع المسار كسلسلة أو None إذا ألغاها المستخدم.
    """
    from pathlib import Path
    from PyQt6.QtWidgets import QInputDialog, QMessageBox
    documents_path = Path('documents')
    documents_path.mkdir(exist_ok=True)
    years = sorted([f.name for f in documents_path.iterdir() if f.is_dir() and f.name.isdigit()], reverse=True)
    year, ok = QInputDialog.getItem(parent, 'اختر السنة', 'السنة:', years + ['سنة جديدة...'], 0, False)
    if not ok:
        return None
    if year == 'سنة جديدة...':
        new_year, ok2 = QInputDialog.getText(parent, 'سنة جديدة', 'أدخل السنة:')
        if ok2 and new_year.isdigit():
            year_folder = documents_path / new_year
            year_folder.mkdir(exist_ok=True)
            QMessageBox.information(parent, 'مجلد السنة', f'تم إنشاء/اختيار مجلد السنة: {year_folder}')
            return str(year_folder)
        else:
            return None
    else:
        year_folder = documents_path / year
        QMessageBox.information(parent, 'مجلد السنة', f'المجلد المختار: {year_folder}')
        return str(year_folder)


class AttachmentDetailsDialog(QDialog):
    """نافذة بسيطة لإدخال معلومات المرفقات مع معاينة الصور"""
    
    def __init__(self, parent=None, scanned_images=[], start_index=0):
        super().__init__(parent)
        self.scanned_images = scanned_images
        self.current_index = start_index
        self.attachment_data = {}
        
        self.setWindowTitle('معلومات المرفقات')
        self.setGeometry(100, 100, 950, 650)
        self.init_ui()
        self.load_attachment(self.current_index)
    
    def init_ui(self):
        main_layout = QVBoxLayout()
        
        # شريط العنوان
        header_layout = QHBoxLayout()
        self.title_label = QLabel()
        self.title_label.setStyleSheet('font-size: 16px; font-weight: bold; color: #2c3e50;')
        header_layout.addWidget(self.title_label)
        header_layout.addStretch()
        
        main_layout.addLayout(header_layout)
        
        # تقسيم النافذة: صورة على اليسار ومعلومات على اليمين
        content_layout = QHBoxLayout()
        
        # قسم الصورة (يسار)
        image_group = QGroupBox('معاينة الصورة')
        image_layout = QVBoxLayout()
        
        self.image_label = QLabel()
        self.image_label.setMinimumSize(380, 480)
        self.image_label.setMaximumSize(380, 480)
        self.image_label.setStyleSheet('border: 2px solid #bdc3c7; background-color: #ecf0f1;')
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.image_label.setScaledContents(False)
        image_layout.addWidget(self.image_label)
        
        image_group.setLayout(image_layout)
        content_layout.addWidget(image_group)
        
        # قسم المعلومات (يمين)
        info_group = QGroupBox('معلومات المرفق')
        info_layout = QVBoxLayout()
        
        form_layout = QFormLayout()
        
        # رقم الوثيقة/المرفق
        self.doc_name = QLineEdit()
        self.doc_name.setPlaceholderText('رقم الوثيقة/المرفق')
        form_layout.addRow('رقم الوثيقة:', self.doc_name)
        
        # التاريخ
        self.doc_date = QLineEdit()
        self.doc_date.setPlaceholderText('مثال: 23-3-2025')
        form_layout.addRow('التاريخ:', self.doc_date)
        
        # المضمون/العنوان
        self.doc_title = QLineEdit()
        self.doc_title.setPlaceholderText('موضوع المرفق')
        form_layout.addRow('المضمون:', self.doc_title)
        
        # جهة الإصدار
        self.issuing_dept = QComboBox()
        self.issuing_dept.addItems(['اختر جهة الإصدار', 'شعبة أمن الأفراد عنة', 'قسم أمن الأفراد الأنبار', 'أخرى'])
        form_layout.addRow('جهة الإصدار:', self.issuing_dept)
        
        # التصنيف
        self.doc_classification = QLineEdit()
        self.doc_classification.setPlaceholderText('التصنيف')
        form_layout.addRow('التصنيف:', self.doc_classification)
        
        # الفقرة القانونية
        self.legal_paragraph = QTextEdit()
        self.legal_paragraph.setMaximumHeight(60)
        self.legal_paragraph.setPlaceholderText('الفقرة القانونية')
        form_layout.addRow('الفقرة القانونية:', self.legal_paragraph)
        
        # ملاحظات خاصة بالمرفق
        self.notes = QTextEdit()
        self.notes.setMaximumHeight(60)
        self.notes.setPlaceholderText('ملاحظات إضافية خاصة بهذا المرفق')
        form_layout.addRow('ملاحظات المرفق:', self.notes)
        
        info_layout.addLayout(form_layout)
        info_group.setLayout(info_layout)
        content_layout.addWidget(info_group)
        
        main_layout.addLayout(content_layout)
        
        # أزرار التنقل
        nav_layout = QHBoxLayout()
        
        self.prev_btn = QPushButton('⏮️ السابق')
        self.prev_btn.clicked.connect(self.go_previous)
        self.prev_btn.setStyleSheet('padding: 8px; font-size: 14px;')
        nav_layout.addWidget(self.prev_btn)
        
        self.position_label = QLabel()
        self.position_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.position_label.setStyleSheet('font-size: 16px; font-weight: bold; color: #2980b9;')
        nav_layout.addWidget(self.position_label)
        
        self.next_btn = QPushButton('التالي ⏭️')
        self.next_btn.clicked.connect(self.go_next)
        self.next_btn.setStyleSheet('padding: 8px; font-size: 14px;')
        nav_layout.addWidget(self.next_btn)
        
        main_layout.addLayout(nav_layout)
        
        # أزرار الإجراءات
        button_layout = QHBoxLayout()
        
        save_all_btn = QPushButton('✅ حفظ والإنهاء')
        save_all_btn.clicked.connect(self.accept)
        save_all_btn.setStyleSheet('background-color: #27ae60; color: white; padding: 10px; font-size: 14px;')
        button_layout.addWidget(save_all_btn)
        
        cancel_btn = QPushButton('❌ إلغاء')
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        main_layout.addLayout(button_layout)
        
        self.setLayout(main_layout)
    
    def load_attachment(self, index):
        """تحميل معلومات وصورة المرفق الحالي"""
        if index < 1 or index >= len(self.scanned_images):
            return
        
        # حفظ البيانات الحالية قبل الانتقال
        if self.current_index >= 1 and self.current_index != index:
            self.save_current_data()
        
        self.current_index = index
        
        # تحديث العنوان
        attachment_num = index  # المرفق الفعلي (0 هو الوثيقة الرئيسية)
        total_attachments = len(self.scanned_images) - 1
        self.title_label.setText(f'📎 المرفق رقم {attachment_num} من {total_attachments}')
        self.position_label.setText(f'{attachment_num} / {total_attachments}')
        
        # تحديث حالة الأزرار
        self.prev_btn.setEnabled(index > 1)
        self.next_btn.setEnabled(index < len(self.scanned_images) - 1)
        
        # تحميل الصورة المصغرة
        try:
            from PIL import Image
            from PyQt6.QtGui import QPixmap
            from PyQt6.QtCore import QTimer
            
            image_path = self.scanned_images[index]
            if os.path.exists(image_path):
                img = Image.open(image_path)
                img.thumbnail((380, 480), Image.Resampling.LANCZOS)
                
                # حفظ مؤقت
                import tempfile
                temp_file = os.path.join(tempfile.gettempdir(), 'preview_temp.jpg')
                img.save(temp_file, 'JPEG')
                
                pixmap = QPixmap(temp_file)
                self.image_label.setPixmap(pixmap)
            else:
                self.image_label.setText('❌ لا يمكن تحميل الصورة')
        except Exception as e:
            self.image_label.setText(f'❌ خطأ في تحميل الصورة:\n{str(e)}')
        
        # تحميل البيانات المحفوظة إن وجدت
        if index in self.attachment_data:
            data = self.attachment_data[index]
            self.doc_name.setText(data.get('doc_name', ''))
            self.doc_date.setText(data.get('doc_date', ''))
            self.doc_title.setText(data.get('doc_title', ''))
            self.doc_classification.setText(data.get('doc_classification', ''))
            self.legal_paragraph.setPlainText(data.get('legal_paragraph', ''))
            self.notes.setPlainText(data.get('notes', ''))
            
            dept = data.get('issuing_dept')
            if dept:
                index_dept = self.issuing_dept.findText(dept)
                if index_dept >= 0:
                    self.issuing_dept.setCurrentIndex(index_dept)
        else:
            # مسح الحقول للمرفق الجديد
            self.doc_name.clear()
            self.doc_date.clear()
            self.doc_title.clear()
            self.doc_classification.clear()
            self.legal_paragraph.clear()
            self.notes.clear()
            self.issuing_dept.setCurrentIndex(0)
    
    def save_current_data(self):
        """حفظ بيانات المرفق الحالي"""
        dept = self.issuing_dept.currentText()
        
        data = {
            'doc_name': self.doc_name.text(),
            'doc_date': self.doc_date.text(),
            'doc_title': self.doc_title.text(),
            'issuing_dept': dept if dept != 'اختر جهة الإصدار' else None,
            'doc_classification': self.doc_classification.text(),
            'legal_paragraph': self.legal_paragraph.toPlainText(),
            'notes': self.notes.toPlainText()
        }
        
        # نتحقق إذا كان هناك أي قيمة غير فارغة
        has_any_data = any(
            v is not None and str(v).strip() != '' 
            for k, v in data.items()
        )
        
        if has_any_data:
            # حفظ البيانات فقط إذا كان هناك معلومات حقيقية
            self.attachment_data[self.current_index] = data
            print(f"[DEBUG] save_current_data: حفظ بيانات المرفق {self.current_index} = {data}")
        else:
            # حذف البيانات إن كانت موجودة (المستخدم مسح كل الحقول)
            if self.current_index in self.attachment_data:
                del self.attachment_data[self.current_index]
                print(f"[DEBUG] save_current_data: حذف بيانات المرفق {self.current_index} (حقول فارغة)")
            else:
                print(f"[DEBUG] save_current_data: تجاهل المرفق {self.current_index} (حقول فارغة)")
    
    def go_previous(self):
        """الانتقال للمرفق السابق"""
        if self.current_index > 1:
            self.save_current_data()
            self.load_attachment(self.current_index - 1)
    
    def go_next(self):
        """الانتقال للمرفق التالي"""
        if self.current_index < len(self.scanned_images) - 1:
            self.save_current_data()
            self.load_attachment(self.current_index + 1)
    
    def accept(self):
        """حفظ جميع البيانات عند الإنهاء"""
        self.save_current_data()
        print(f"[DEBUG] accept: جميع البيانات المحفوظة = {self.attachment_data}")
        super().accept()
    
    def get_all_data(self):
        """الحصول على بيانات جميع المرفقات"""
        return self.attachment_data


class AddDocumentDialog(QDialog):
    def select_year_folder(self):
        """استخدم الدالة المساعدة الموحدة لاختيار مجلد السنة."""
        try:
            return choose_year_folder(self)
        except Exception:
            return None

    """نافذة حوار لإضافة وثيقة جديدة"""
    
    def __init__(self, parent=None, db=None, image_manager=None):
        super().__init__(parent)
        self.setWindowTitle('إضافة وثيقة جديدة')
        self.setGeometry(100, 100, 600, 500)
        self.db = db
        self.image_manager = image_manager
        self.scanned_image_path = None
        self.scanned_images = []  # قائمة لحفظ عدة صور
        self.attachment_details_dict = {}  # قاموس لحفظ بيانات المرفقات
        self.init_ui()
    
    def init_ui(self):
        layout = QFormLayout()
        
        # اسم الوثيقة
        self.doc_name = QLineEdit()
        layout.addRow('اسم الوثيقة:', self.doc_name)
        
        # التاريخ
        self.doc_date = QLineEdit()
        self.doc_date.setPlaceholderText('مثال: 23-3-2025')
        layout.addRow('تاريخ الوثيقة:', self.doc_date)
        
        # العنوان
        self.doc_title = QLineEdit()
        layout.addRow('عنوان الوثيقة:', self.doc_title)
        
        # جهة الإصدار
        self.issuing_dept = QComboBox()
        self.issuing_dept.addItems(['اختر جهة الإصدار', 'شعبة أمن الأفراد عنة', 'قسم أمن الأفراد الأنبار'])
        layout.addRow('جهة الإصدار:', self.issuing_dept)
        
        # تصنيف الوثيقة
        self.doc_classification = QLineEdit()
        layout.addRow('تصنيف الوثيقة:', self.doc_classification)
        
        # الفقرة القانونية
        self.legal_paragraph = QTextEdit()
        self.legal_paragraph.setMaximumHeight(100)
        layout.addRow('الفقرة القانونية:', self.legal_paragraph)
        
        # عدد الوجوه
        self.sides = QSpinBox()
        self.sides.setMinimum(1)
        self.sides.setMaximum(2)
        self.sides.setValue(1)
        layout.addRow('عدد الوجوه:', self.sides)
        
        # قائمة الصور الممسوحة
        self.images_label = QLabel('عدد الصور الممسوحة: 0')
        layout.addRow(self.images_label)
        
        # حالة السكانر (Disclaimer)
        self.scanner_status_label = QLabel()
        self._update_scanner_status()
        layout.addRow(self.scanner_status_label)

        # مجلد السنة (اختيار مجلد السنة داخل مجلد التخزين 'documents')
        from PyQt6.QtWidgets import QWidget
        year_widget = QWidget()
        year_widget_layout = QHBoxLayout()
        self.year_folder_edit = QLineEdit()
        self.year_folder_edit.setReadOnly(True)
        self.year_folder_edit.setPlaceholderText('لم يتم اختيار مجلد السنة')
        year_select_btn = QPushButton('📂 اختيار مجلد السنة')
        year_select_btn.clicked.connect(self.on_choose_year_folder)
        year_widget_layout.addWidget(self.year_folder_edit)
        year_widget_layout.addWidget(year_select_btn)
        year_widget.setLayout(year_widget_layout)
        layout.addRow('مجلد السنة:', year_widget)
        
        # أزرار المسح
        scan_layout = QHBoxLayout()
        
        scan_one_btn = QPushButton('📷 مسح صورة واحدة')
        scan_one_btn.clicked.connect(self.scan_manual)
        scan_layout.addWidget(scan_one_btn)
        
        scan_multiple_btn = QPushButton('📚 مسح مرفقات متعددة')
        scan_multiple_btn.clicked.connect(self.scan_multiple)
        scan_layout.addWidget(scan_multiple_btn)
        
        layout.addRow(scan_layout)
        
        ocr_btn = QPushButton('🔍 استخراج تلقائي من الصور (بطيء)')
        ocr_btn.clicked.connect(self.scan_and_extract)
        layout.addRow(ocr_btn)
        
        # أزرار
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addRow(button_box)
        
        self.setLayout(layout)
    
    def _update_scanner_status(self):
        """تحديث حالة السكانر في الواجهة"""
        global SCANNER_AVAILABLE, SCANNER_COUNT
        
        # إعادة التحقق من السكانرات المتصلة
        if SCANNER_AVAILABLE:
            try:
                _wia_manager = win32com.client.Dispatch("WIA.DeviceManager")
                SCANNER_COUNT = _wia_manager.DeviceInfos.Count
            except:
                SCANNER_COUNT = 0
        
        if not SCANNER_AVAILABLE:
            self.scanner_status_label.setText('⚠️ حالة السكانر: مكتبة pywin32 غير مثبتة - استخدم اختيار الصور من الحاسب')
            self.scanner_status_label.setStyleSheet('color: #e74c3c; font-size: 11px; padding: 5px; background-color: #fdf2f2; border-radius: 3px;')
        elif SCANNER_COUNT == 0:
            self.scanner_status_label.setText('⚠️ حالة السكانر: لا يوجد سكانر متصل - قم بتوصيل السكانر أو اختر صورة من الحاسب')
            self.scanner_status_label.setStyleSheet('color: #e67e22; font-size: 11px; padding: 5px; background-color: #fef9e7; border-radius: 3px;')
        else:
            self.scanner_status_label.setText(f'✅ حالة السكانر: متصل ({SCANNER_COUNT} جهاز)')
            self.scanner_status_label.setStyleSheet('color: #27ae60; font-size: 11px; padding: 5px; background-color: #eafaf1; border-radius: 3px;')

    def on_choose_year_folder(self):
        """مستدعى عند الضغط على زر اختيار مجلد السنة"""
        path = self.select_year_folder()
        if path:
            self.selected_year_folder = path
            try:
                self.year_folder_edit.setText(path)
            except Exception:
                pass
    
    def scan_manual(self):
        """مسح من السكانر مع إدخال يدوي (سريع)"""
        global SCANNER_AVAILABLE, SCANNER_COUNT
        
        # التحقق من توفر المكتبة
        if not SCANNER_AVAILABLE:
            reply = QMessageBox.question(
                self, 'السكانر غير متاح',
                '⚠️ مكتبة السكانر (pywin32) غير مثبتة\n\n'
                'لتثبيتها، قم بتشغيل الأمر التالي:\n'
                'pip install pywin32\n\n'
                'هل تريد اختيار صورة من الحاسب بدلاً من المسح؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._select_image_file()
            return
        
        # التحقق من وجود سكانر متصل
        self._update_scanner_status()
        if SCANNER_COUNT == 0:
            reply = QMessageBox.question(
                self, 'السكانر غير متصل',
                '⚠️ لا يوجد سكانر متصل بالحاسب\n\n'
                'تأكد من:\n'
                '• توصيل السكانر بالحاسب\n'
                '• تشغيل السكانر\n'
                '• تثبيت برنامج تشغيل السكانر\n\n'
                'هل تريد اختيار صورة من الحاسب بدلاً من المسح؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._select_image_file()
            return
        
        try:
            # استخدم اختيار مجلد السنة المدمج إن كان موجوداً، وإلا اطلبه عبر الحوار الموحد
            year_folder = getattr(self, 'selected_year_folder', None)
            if not year_folder:
                try:
                    txt = self.year_folder_edit.text().strip()
                    if txt and txt != 'لم يتم اختيار مجلد السنة':
                        year_folder = txt
                except Exception:
                    year_folder = None

            if not year_folder:
                year_folder = self.select_year_folder()
                if not year_folder:
                    return
            self.selected_year_folder = year_folder
            try:
                self.year_folder_edit.setText(year_folder)
            except Exception:
                pass

            QMessageBox.information(
                self, 'جاري المسح',
                'سيتم فتح نافذة السكانر\n\nضع الوثيقة واضغط Scan'
            )
            
            # فتح نافذة السكانر
            wia = win32com.client.Dispatch("WIA.CommonDialog")
            image = wia.ShowAcquireImage()
            
            if not image:
                return
            
            # إنشاء اسم ملف فريد
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f'scanned_{timestamp}.jpg')
            
            # حذف الملف إذا كان موجوداً
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass
            
            # حفظ الصورة
            image.SaveFile(temp_file)
            
            # حفظ مسار الصورة
            self.scanned_image_path = temp_file
            self.scanned_images = [temp_file]  # إضافة للقائمة أيضاً
            self._update_images_count()
            
            QMessageBox.information(
                self, 'تم المسح ✅',
                'تم مسح الوثيقة بنجاح!\n\nأدخل المعلومات يدوياً في الحقول أدناه'
            )
            
        except Exception as e:
            QMessageBox.critical(
                self, 'خطأ',
                f'خطأ في المسح الضوئي:\n{str(e)}\n\nتأكد من:\n• توصيل السكانر\n• تثبيت برنامج السكانر'
            )
    
    def _select_image_file(self):
        """اختيار صورة من الحاسب كبديل للسكانر مع اختيار مجلد السنة"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, 'اختر صورة الوثيقة',
            '', 'صور (*.jpg *.jpeg *.png *.tiff *.bmp);;جميع الملفات (*)'
        )
        if file_path:
            # استخدم اختيار مجلد السنة المدمج إذا اختاره المستخدم مسبقاً
            year_folder = getattr(self, 'selected_year_folder', None)
            if not year_folder:
                try:
                    txt = self.year_folder_edit.text().strip()
                    if txt and txt != 'لم يتم اختيار مجلد السنة':
                        year_folder = txt
                except Exception:
                    year_folder = None

            if not year_folder:
                year_folder = self.select_year_folder()
                if not year_folder:
                    QMessageBox.warning(self, 'تنبيه', 'يجب اختيار أو إنشاء مجلد سنة')
                    return

            import os, shutil
            basename = os.path.basename(file_path)
            dest_path = os.path.join(year_folder, basename)
            shutil.copy2(file_path, dest_path)
            self.scanned_image_path = dest_path
            self.scanned_images = [dest_path]
            self._update_images_count()
            QMessageBox.information(
                self, 'تم ✅',
                f'تم اختيار الصورة بنجاح!\n\nتم نقلها إلى مجلد السنة: {year_folder}\nأدخل المعلومات يدوياً في الحقول أدناه'
            )
    
    def _select_multiple_image_files(self):
        """اختيار عدة صور من الحاسب كبديل للسكانر"""
        files, _ = QFileDialog.getOpenFileNames(
            self, 'اختر صور الوثائق',
            '', 'صور (*.jpg *.jpeg *.png *.tiff *.bmp);;جميع الملفات (*)'
        )
        
        if files:
            # تأكد من اختيار مجلد السنة (استخدم الحقل المدمج أولاً)
            year_folder = getattr(self, 'selected_year_folder', None)
            if not year_folder:
                try:
                    txt = self.year_folder_edit.text().strip()
                    if txt and txt != 'لم يتم اختيار مجلد السنة':
                        year_folder = txt
                except Exception:
                    year_folder = None

            if not year_folder:
                year_folder = self.select_year_folder()
                if not year_folder:
                    QMessageBox.warning(self, 'تنبيه', 'يجب اختيار أو إنشاء مجلد سنة')
                    return
            self.selected_year_folder = year_folder
            try:
                self.year_folder_edit.setText(year_folder)
            except Exception:
                pass

            # انسخ الملفات إلى مجلد السنة حتى تعمل عمليات الحفظ لاحقاً بنفس المسار
            import shutil, os
            dest_files = []
            for f in files:
                try:
                    basename = os.path.basename(f)
                    dest = os.path.join(year_folder, basename)
                    # تجنب الكتابة فوق ملفات موجودة بإضافة طابع زمني إن لزم
                    if os.path.exists(dest):
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
                        name, ext = os.path.splitext(basename)
                        dest = os.path.join(year_folder, f"{name}_{timestamp}{ext}")
                    shutil.copy2(f, dest)
                    dest_files.append(dest)
                except Exception as e:
                    print(f"خطأ في نسخ الملف {f}: {e}")

            if dest_files:
                self.scanned_images = dest_files
                self.scanned_image_path = dest_files[0] if dest_files else None
                self._update_images_count()

                if len(dest_files) > 1:
                    self._handle_scanned_documents(len(dest_files))
                else:
                    QMessageBox.information(
                        self, 'تم ✅',
                        'تم اختيار الصورة ونقلها لمجلد السنة بنجاح!\n\nأدخل المعلومات يدوياً في الحقول أدناه'
                    )
            else:
                QMessageBox.warning(self, 'تنبيه', 'لم يتم نسخ أي ملفات')
    
    def scan_multiple(self):
        """مسح تلقائي لجميع الأوراق دفعة واحدة"""
        global SCANNER_AVAILABLE, SCANNER_COUNT
        
        # التحقق من توفر المكتبة
        if not SCANNER_AVAILABLE:
            reply = QMessageBox.question(
                self, 'السكانر غير متاح',
                '⚠️ مكتبة السكانر (pywin32) غير مثبتة\n\n'
                'لتثبيتها، قم بتشغيل الأمر التالي:\n'
                'pip install pywin32\n\n'
                'هل تريد اختيار عدة صور من الحاسب بدلاً من المسح؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._select_multiple_image_files()
            return
        
        # التحقق من وجود سكانر متصل
        self._update_scanner_status()
        if SCANNER_COUNT == 0:
            reply = QMessageBox.question(
                self, 'السكانر غير متصل',
                '⚠️ لا يوجد سكانر متصل بالحاسب\n\n'
                'تأكد من:\n'
                '• توصيل السكانر بالحاسب\n'
                '• تشغيل السكانر\n'
                '• تثبيت برنامج تشغيل السكانر\n\n'
                'هل تريد اختيار عدة صور من الحاسب بدلاً من المسح؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._select_multiple_image_files()
            return
        
        try:
            reply = QMessageBox.question(
                self, 'مسح تلقائي جماعي',
                '🔄 مسح تلقائي مستمر لجميع الأوراق\n\n'
                '✅ ضع جميع الأوراق في وحدة التغذية (ADF)\n'
                '✅ أو رتبها جاهزة للمسح المتتالي\n\n'
                'سيتم مسح جميع الأوراق تلقائياً بدون توقف\n\n'
                'هل تريد البدء؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply != QMessageBox.StandardButton.Yes:
                return

            # تأكد من وجود مجلد السنة المختار (الحقل المدمج أولاً)
            year_folder = getattr(self, 'selected_year_folder', None)
            if not year_folder:
                try:
                    txt = self.year_folder_edit.text().strip()
                    if txt and txt != 'لم يتم اختيار مجلد السنة':
                        year_folder = txt
                except Exception:
                    year_folder = None

            if not year_folder:
                year_folder = self.select_year_folder()
                if not year_folder:
                    QMessageBox.information(self, 'ملغى', 'تم إلغاء المسح - يجب اختيار مجلد السنة للمتابعة')
                    return
            self.selected_year_folder = year_folder
            try:
                self.year_folder_edit.setText(year_folder)
            except Exception:
                pass
            
            # محاولة استخدام المسح التلقائي أولاً
            try:
                scan_count = self._scan_automatic_feeder()
                if scan_count > 0:
                    self._handle_scanned_documents(scan_count)
                    return
            except Exception as e:
                print(f"المسح التلقائي فشل: {e}")
            
            # إذا فشل المسح التلقائي، استخدم المسح المتتالي السريع
            self._scan_continuous_manual()
            
        except Exception as e:
            QMessageBox.critical(
                self, 'خطأ',
                f'خطأ في المسح الضوئي:\n{str(e)}\n\nتأكد من:\n• توصيل السكانر\n• تثبيت برنامج السكانر'
            )
    
    def _scan_automatic_feeder(self):
        """مسح تلقائي باستخدام وحدة التغذية التلقائية (ADF) مع إمكانية استئناف"""
        wia = win32com.client.Dispatch("WIA.DeviceManager")
        
        if wia.DeviceInfos.Count == 0:
            raise Exception("لا يوجد سكانر متصل")
        
        device_info = wia.DeviceInfos.Item(1)
        device = device_info.Connect()
        
        # تفعيل وحدة التغذية التلقائية
        try:
            for prop in device.Properties:
                if "Document Handling Select" in str(prop.Name):
                    prop.Value = 1  # FEEDER
                    break
        except:
            pass
        
        scan_count = len(self.scanned_images)  # البدء من آخر عدد ممسوح
        temp_dir = tempfile.gettempdir()
        
        QMessageBox.information(
            self, 'جاري المسح...',
            '⏳ المسح التلقائي بدأ\n\n'
            'سيتم مسح جميع الأوراق تلقائياً\n'
            'الرجاء الانتظار...'
        )
        
        # مسح مستمر حتى انتهاء الأوراق
        while True:
            try:
                # محاولة مسح الورقة التالية
                item = device.Items[1]
                image = item.Transfer("{B96B3CAE-0728-11D3-9D7B-0000F81EF32E}")
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                temp_file = os.path.join(temp_dir, f'auto_scan_{scan_count+1}_{timestamp}.jpg')
                
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                
                image.SaveFile(temp_file)
                self.scanned_images.append(temp_file)
                scan_count += 1
                self._update_images_count()
                
            except Exception as e:
                # التحقق من وجود صور ممسوحة
                if scan_count > 0:
                    # عرض خيار الاستئناف
                    reply = QMessageBox.question(
                        self, '⚠️ انتهى المسح أو حدث خطأ',
                        f'✅ تم مسح {scan_count} ورقة\n\n'
                        'انتهت الأوراق أو حدث خطأ في السكانر\n\n'
                        'هل تريد إضافة المزيد من الأوراق والاستمرار؟\n\n'
                        '• نعم: أضف أوراق واستمر\n'
                        '• لا: إنهاء المسح',
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    
                    if reply == QMessageBox.StandardButton.Yes:
                        QMessageBox.information(
                            self, 'استئناف المسح',
                            'أضف الأوراق الإضافية إلى وحدة التغذية\n\n'
                            'اضغط OK للمتابعة'
                        )
                        
                        try:
                            # إعادة الاتصال بالجهاز
                            device = device_info.Connect()
                            continue
                        except:
                            QMessageBox.warning(
                                self, 'تحذير',
                                'فشل إعادة الاتصال\nسيتم حفظ ما تم مسحه'
                            )
                            break
                    else:
                        break
                else:
                    # لم يتم مسح أي شيء، رفع الخطأ
                    break
        
        return scan_count
    
    def _scan_continuous_manual(self):
        """مسح متتالي سريع بدون نوافذ متكررة مع إمكانية استئناف"""
        QMessageBox.information(
            self, 'مسح متتالي',
            '📄 مسح متتالي سريع\n\n'
            'ضع جميع الأوراق جاهزة\n'
            'سيتم فتح نافذة السكانر مباشرة\n\n'
            'امسح جميع الأوراق واحدة تلو الأخرى\n'
            'اضغط Cancel عند الانتهاء من آخر ورقة\n\n'
            '⚠️ في حالة حدوث خطأ، يمكنك استئناف المسح'
        )
        
        wia = win32com.client.Dispatch("WIA.CommonDialog")
        scan_count = len(self.scanned_images)  # البدء من آخر عدد ممسوح
        temp_dir = tempfile.gettempdir()
        consecutive_errors = 0
        
        # مسح مستمر بدون نوافذ متكررة
        while True:
            try:
                # فتح نافذة السكانر مباشرة
                image = wia.ShowAcquireImage()
                
                if not image:
                    break
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                temp_file = os.path.join(temp_dir, f'scanned_{scan_count+1}_{timestamp}.jpg')
                
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                    except:
                        pass
                
                image.SaveFile(temp_file)
                self.scanned_images.append(temp_file)
                scan_count += 1
                self._update_images_count()
                consecutive_errors = 0  # إعادة تعيين عداد الأخطاء
                
            except Exception as e:
                consecutive_errors += 1
                
                # إذا كان هناك صور ممسوحة، اعرض خيار الاستئناف
                if scan_count > 0:
                    error_msg = str(e)
                    reply = QMessageBox.question(
                        self, '⚠️ خطأ في السكانر',
                        f'حدث خطأ في السكانر:\n{error_msg}\n\n'
                        f'✅ تم مسح {scan_count} ورقة بنجاح\n\n'
                        'ماذا تريد أن تفعل؟\n\n'
                        '• نعم: حل المشكلة واستئناف المسح\n'
                        '• لا: إنهاء وحفظ ما تم مسحه',
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    
                    if reply == QMessageBox.StandardButton.Yes:
                        # استئناف المسح
                        QMessageBox.information(
                            self, 'استئناف المسح',
                            '🔧 قم بمعالجة مشكلة السكانر:\n\n'
                            '1. تحقق من توصيل السكانر\n'
                            '2. تحقق من وجود ورق في السكانر\n'
                            '3. أعد تشغيل السكانر إن لزم الأمر\n'
                            '4. اضغط OK للمتابعة'
                        )
                        
                        # محاولة إعادة الاتصال بالسكانر
                        try:
                            wia = win32com.client.Dispatch("WIA.CommonDialog")
                            consecutive_errors = 0
                            continue  # استئناف الحلقة
                        except:
                            QMessageBox.warning(
                                self, 'تحذير',
                                'فشل إعادة الاتصال بالسكانر\nسيتم حفظ ما تم مسحه'
                            )
                            break
                    else:
                        # إنهاء وحفظ ما تم مسحه
                        break
                else:
                    # لم يتم مسح أي شيء بعد
                    if consecutive_errors >= 2:
                        raise e
                    
                    reply = QMessageBox.question(
                        self, 'خطأ في السكانر',
                        f'حدث خطأ في السكانر:\n{str(e)}\n\n'
                        'هل تريد المحاولة مرة أخرى؟',
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    
                    if reply == QMessageBox.StandardButton.Yes:
                        try:
                            wia = win32com.client.Dispatch("WIA.CommonDialog")
                            continue
                        except:
                            raise e
                    else:
                        raise e
        
        if scan_count > 0:
            self._handle_scanned_documents(scan_count)
    
    def _handle_scanned_documents(self, scan_count):
        """معالجة الوثائق الممسوحة"""
        self.scanned_image_path = self.scanned_images[0]
        
        # سؤال عن معلومات المرفقات
        if scan_count > 1:
            reply = QMessageBox.question(
                self, 'معلومات المرفقات',
                f'✅ تم مسح {scan_count} وثيقة/مرفق بنجاح!\n\n'
                f'• الورقة الأولى: الوثيقة الرئيسية\n'
                f'• {scan_count - 1} ورقة: مرفقات\n\n'
                'هل تريد إدخال معلومات منفصلة لكل مرفق؟\n\n'
                '• نعم: إدخال معلومات لكل مرفق\n'
                '• لا: حفظ الكل مع الوثيقة الرئيسية فقط',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self._collect_attachment_details()
            else:
                QMessageBox.information(
                    self, 'تم ✅',
                    f'تم مسح {scan_count} صفحة (1 رئيسية + {scan_count-1} مرفق)\n\n'
                    'أدخل معلومات الوثيقة الرئيسية في الحقول أدناه\n'
                    'سيتم حفظ جميع المرفقات معها'
                )
        else:
            QMessageBox.information(
                self, 'تم المسح ✅',
                'تم مسح الوثيقة الرئيسية\n\nأدخل المعلومات في الحقول أدناه'
            )
    
    def _collect_attachment_details(self):
        """جمع معلومات تفصيلية لجميع المرفقات باستخدام نافذة واحدة مع التنقل"""
        
        print(f"[DEBUG] _collect_attachment_details: بدء - عدد الصور = {len(self.scanned_images)}")
        
        # فتح نافذة واحدة لجميع المرفقات مع إمكانية التنقل
        # start_index=1 يعني أن المرفق الأول (الصورة الثانية) يبدأ من الفهرس 1
        dialog = AttachmentDetailsDialog(
            self, 
            self.scanned_images, 
            start_index=1
        )
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            # الحصول على بيانات جميع المرفقات
            # all_data هو قاموس بمفاتيح 1, 2, 3, ... 
            # حيث المفتاح 1 = المرفق الأول (الصورة الثانية في قائمة scanned_images)
            all_data = dialog.get_all_data()
            
            print(f"[DEBUG] _collect_attachment_details: all_data من الحوار = {all_data}")
            print(f"[DEBUG] _collect_attachment_details: نوع all_data = {type(all_data)}")
            print(f"[DEBUG] _collect_attachment_details: مفاتيح all_data = {list(all_data.keys())}")
            
            # نحفظ القاموس مباشرة بدون تحويل لقائمة
            # المفتاح في all_data يتوافق مع فهرس الصورة في scanned_images
            # المفتاح 1 = scanned_images[1] = الصورة الثانية = المرفق الأول
            self.attachment_details_dict = dict(all_data)  # نسخة عميقة
            
            print(f"[DEBUG] _collect_attachment_details: تم الحفظ في self.attachment_details_dict = {self.attachment_details_dict}")
            
            valid_count = len([v for v in all_data.values() if v and any(str(x).strip() for x in v.values() if x)])
            QMessageBox.information(
                self, 'تم ✅',
                f'تم إدخال معلومات {valid_count} مرفق بنجاح\n\n'
                f'[تصحيح] البيانات المحفوظة: {len(self.attachment_details_dict)} مرفق\n\n'
                'أدخل معلومات الوثيقة الرئيسية في الحقول أدناه'
            )
        else:
            # تم الإلغاء
            self.attachment_details_dict = {}
            print(f"[DEBUG] _collect_attachment_details: تم الإلغاء")
            QMessageBox.information(
                self, 'تم',
                'سيتم استخدام معلومات الوثيقة الرئيسية لجميع المرفقات'
            )
    
    def _update_images_count(self):
        """تحديث عدد الصور الممسوحة"""
        count = len(self.scanned_images)
        self.images_label.setText(f'عدد الصور الممسوحة: {count}')
    
    def scan_and_extract(self):
        """مسح من السكانر واستخراج المعلومات تلقائياً (بطيء)"""
        global SCANNER_AVAILABLE, SCANNER_COUNT
        
        # التحقق من توفر المكتبة
        if not SCANNER_AVAILABLE:
            reply = QMessageBox.question(
                self, 'السكانر غير متاح',
                '⚠️ مكتبة السكانر (pywin32) غير مثبتة\n\n'
                'لتثبيتها، قم بتشغيل الأمر التالي:\n'
                'pip install pywin32\n\n'
                'هل تريد اختيار صورة من الحاسب لاستخراج المعلومات منها؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.extract_from_image()
            return
        
        # التحقق من وجود سكانر متصل
        self._update_scanner_status()
        if SCANNER_COUNT == 0:
            reply = QMessageBox.question(
                self, 'السكانر غير متصل',
                '⚠️ لا يوجد سكانر متصل بالحاسب\n\n'
                'تأكد من:\n'
                '• توصيل السكانر بالحاسب\n'
                '• تشغيل السكانر\n'
                '• تثبيت برنامج تشغيل السكانر\n\n'
                'هل تريد اختيار صورة من الحاسب لاستخراج المعلومات منها؟',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.extract_from_image()
            return
        
        try:
            reply = QMessageBox.warning(
                self, 'تحذير',
                'الاستخراج التلقائي قد يستغرق 1-2 دقيقة!\n\n'
                'هل تريد المتابعة؟\n\n'
                'للمسح السريع: استخدم "مسح من السكانر (إدخال يدوي)"',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply != QMessageBox.StandardButton.Yes:
                return

            # اطلب من المستخدم اختيار مجلد السنة قبل المسح
            # استخدم اختيار مجلد السنة المدمج إن كان موجوداً، وإلا اطلبه عبر الحوار الموحد
            year_folder = getattr(self, 'selected_year_folder', None)
            if not year_folder:
                try:
                    txt = self.year_folder_edit.text().strip()
                    if txt and txt != 'لم يتم اختيار مجلد السنة':
                        year_folder = txt
                except Exception:
                    year_folder = None

            if not year_folder:
                year_folder = self.select_year_folder()
                if not year_folder:
                    return
            self.selected_year_folder = year_folder
            try:
                self.year_folder_edit.setText(year_folder)
            except Exception:
                pass

            QMessageBox.information(
                self, 'جاري المسح',
                'سيتم فتح نافذة السكانر\n\nضع الوثيقة في السكانر واضغط Scan'
            )
            
            # فتح نافذة السكانر
            wia = win32com.client.Dispatch("WIA.CommonDialog")
            image = wia.ShowAcquireImage()
            
            if not image:
                return
            
            # إنشاء اسم ملف فريد باستخدام الوقت
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f'scanned_{timestamp}.jpg')
            
            # حذف الملف إذا كان موجوداً
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass
            
            # حفظ الصورة
            image.SaveFile(temp_file)
            
            # التحقق من وجود الملف
            if not os.path.exists(temp_file):
                raise Exception("فشل حفظ الصورة الممسوحة")
            
            # استخراج المعلومات (بطيء)
            self._process_scanned_image(temp_file)
            
        except Exception as e:
            QMessageBox.critical(
                self, 'خطأ',
                f'خطأ في المسح الضوئي:\n{str(e)}\n\nتأكد من:\n• توصيل السكانر\n• تثبيت برنامج السكانر'
            )
    
    def _process_scanned_image(self, image_path):
        """معالجة الصورة الممسوحة واستخراج المعلومات وحفظها تلقائياً في خيط خلفي"""
        from PyQt6.QtCore import QThread, pyqtSignal, QObject

        class OCRWorker(QThread):
            finished = pyqtSignal(object)
            failed = pyqtSignal(str)

            def __init__(self, img_path):
                super().__init__()
                self.img_path = img_path

            def run(self):
                try:
                    from app.ocr_extractor import OCRExtractor
                    extractor = OCRExtractor()
                    info = extractor.extract_document_info(self.img_path)
                    self.finished.emit(info)
                except Exception as e:
                    self.failed.emit(str(e))

        # إظهار مربع تقدم بسيط وقابل للإلغاء
        progress = QProgressDialog('جاري استخراج المعلومات...', 'إلغاء', 0, 0, self)
        progress.setWindowTitle('استخراج المضمون')
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.show()

        worker = OCRWorker(image_path)

        def on_finished(info):
            progress.close()
            try:
                if info and (info.get('doc_number') or info.get('doc_date')):
                    self._fill_fields(info)
                    self._save_with_image(image_path, info)
                else:
                    QMessageBox.warning(self, 'تنبيه', 'لم يتم استخراج معلومات. أدخلها يدوياً')
                    self.scanned_image_path = image_path
            except Exception as e:
                QMessageBox.critical(self, 'خطأ', f'خطأ في معالجة نتيجة الاستخراج: {str(e)}')

        def on_failed(err):
            progress.close()
            QMessageBox.critical(self, 'خطأ', f'خطأ في الاستخراج: {err}\n\nأدخل المعلومات يدوياً')
            self.scanned_image_path = image_path

        worker.finished.connect(on_finished)
        worker.failed.connect(on_failed)

        # احتفظ بالمرجع لمنع جمع القمامة
        self._ocr_worker = worker
        worker.start()
    
    def _fill_fields(self, info):
        """ملء الحقول بالمعلومات المستخرجة"""
        if info['doc_number']:
            self.doc_name.setText(info['doc_number'])
        if info['doc_date']:
            self.doc_date.setText(info['doc_date'])
        if info['doc_title']:
            self.doc_title.setText(info['doc_title'])
        if info['issuing_dept']:
            index = self.issuing_dept.findText(info['issuing_dept'])
            if index >= 0:
                self.issuing_dept.setCurrentIndex(index)
    
    def _save_with_image(self, image_path, info):
        """حفظ الوثيقة والصورة تلقائياً"""
        if not self.db or not self.image_manager:
            return
        
        try:
            doc_name = info['doc_number']
            if info['doc_date']:
                doc_name += f" في {info['doc_date']}"
            
            doc_id = self.db.add_document(
                doc_name,
                info['doc_date'],
                info['doc_title'],
                info['issuing_dept'],
                '', ''
            )
            
            if image_path and os.path.exists(image_path):
                # استخدم مجلد السنة المختار مسبقاً إذا وُجد، وإلا اطلبه الآن
                year_folder = getattr(self, 'selected_year_folder', None)
                if not year_folder:
                    year_folder = self.select_year_folder()
                if not year_folder:
                    QMessageBox.warning(self, 'تنبيه', 'يجب اختيار أو إنشاء مجلد سنة لحفظ الصورة')
                    return

                # انسخ الصورة الممسوحة إلى مجلد مؤقت داخل مجلد السنة حتى يتعرف ImageManager على السنة
                import shutil, os
                incoming_dir = os.path.join(year_folder, '_incoming')
                os.makedirs(incoming_dir, exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
                basename = os.path.basename(image_path)
                temp_name = f"{timestamp}_{basename}"
                temp_dest = os.path.join(incoming_dir, temp_name)
                shutil.copy2(image_path, temp_dest)

                # احفظ الصورة باستخدام ImageManager (سيضعها ضمن مجلد السنة تلقائياً)
                # مرّر اسم السنة إلى ImageManager لتأكيد الحفظ داخل مجلد السنة
                try:
                    year_name = Path(year_folder).name
                except Exception:
                    year_name = None
                saved_path = self.image_manager.save_image(temp_dest, doc_id, 1, year=year_name)
                # أضف إلى قاعدة البيانات
                self.db.add_image(doc_id, saved_path, basename, 1, None, 1, None)

                # حذف الملف المؤقت في incoming
                try:
                    os.remove(temp_dest)
                except Exception:
                    pass
            
            QMessageBox.information(
                self, 'تم الحفظ ✅',
                f'تم حفظ الوثيقة:\n• {doc_name}\n\nيمكنك البحث عنها وطباعتها'
            )
            self.accept()
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'خطأ في الحفظ: {str(e)}')
    
    def extract_from_image(self):
        """استخراج المعلومات من ملف صورة موجود"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, 'اختر صورة الوثيقة',
            '', 'صور (*.jpg *.jpeg *.png *.tiff *.bmp);;جميع الملفات (*)'
        )
        
        if file_path:
            self._process_scanned_image(file_path)
    
    def get_data(self):
        """الحصول على البيانات المدخلة"""
        dept = self.issuing_dept.currentText()
        if dept == 'اختر جهة الإصدار':
            dept = ''
        
        # التأكد من أن attachment_details_dict موجود
        att_dict = getattr(self, 'attachment_details_dict', {})
        print(f"[DEBUG] get_data: attachment_details_dict = {att_dict}")
        print(f"[DEBUG] get_data: نوعه = {type(att_dict)}")
        print(f"[DEBUG] get_data: مفاتيحه = {list(att_dict.keys()) if att_dict else 'فارغ'}")
        
        return {
            'doc_name': self.doc_name.text(),
            'doc_date': self.doc_date.text(),
            'doc_title': self.doc_title.text(),
            'issuing_dept': dept,
            'doc_classification': self.doc_classification.text(),
            'legal_paragraph': self.legal_paragraph.toPlainText(),
            'sides': self.sides.value(),
            'scanned_image': self.scanned_image_path,
            'scanned_images': self.scanned_images,
            'attachment_details_dict': att_dict,
            'selected_year_folder': getattr(self, 'selected_year_folder', None)
        }


class ImportImagesDialog(QDialog):
    """نافذة حوار لاستيراد الصور"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('استيراد الصور')
        self.setGeometry(100, 100, 700, 500)
        self.selected_files = []
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # أزرار الاستيراد
        button_layout = QHBoxLayout()
        
        select_btn = QPushButton('📄 اختر صور من الحاسب')
        select_btn.clicked.connect(self.select_files)
        button_layout.addWidget(select_btn)
        
        select_folder_btn = QPushButton('📁 اختر مجلد كامل')
        select_folder_btn.clicked.connect(self.select_folder)
        button_layout.addWidget(select_folder_btn)
        
        layout.addLayout(button_layout)
        
        # قائمة الملفات المختارة
        file_label_layout = QHBoxLayout()
        file_label_layout.addWidget(QLabel('الملفات المختارة:'))
        file_label_layout.addStretch()
        
        select_all_btn = QPushButton('✓ تحديد الكل')
        select_all_btn.clicked.connect(self.select_all_files)
        select_all_btn.setMaximumWidth(100)
        file_label_layout.addWidget(select_all_btn)
        
        delete_btn = QPushButton('🗑️ حذف المحددة')
        delete_btn.clicked.connect(self.delete_selected)
        delete_btn.setMaximumWidth(100)
        file_label_layout.addWidget(delete_btn)
        
        layout.addLayout(file_label_layout)
        
        self.file_list = QListWidget()
        self.file_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        layout.addWidget(self.file_list)
        
        # معلومات تحليل
        layout.addWidget(QLabel('معلومات الملفات المكتشفة:'))
        self.info_text = QTextEdit()
        self.info_text.setReadOnly(True)
        self.info_text.setMaximumHeight(150)
        layout.addWidget(self.info_text)
        
        # أزرار
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
    
    def select_files(self):
        """اختيار الملفات"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            'اختر صور الوثائق',
            '',
            'صور (*.jpg *.jpeg *.png *.tiff *.bmp)'
        )
        
        if files:
            # اطلب من المستخدم اختيار مجلد السنة لحفظ الصور المختارة
            year_folder = choose_year_folder(self)
            if not year_folder:
                QMessageBox.warning(self, 'تنبيه', 'يجب اختيار أو إنشاء مجلد سنة')
                return

            import os, shutil
            dest_files = []
            for f in files:
                try:
                    basename = os.path.basename(f)
                    dest = os.path.join(year_folder, basename)
                    shutil.copy2(f, dest)
                    dest_files.append(dest)
                except Exception as e:
                    print(f"خطأ في نسخ الملف {f}: {e}")

            if dest_files:
                self.selected_files = dest_files
                self.update_list()
            else:
                QMessageBox.warning(self, 'تنبيه', 'لم يتم نسخ أي ملفات')
    
    def select_folder(self):
        """اختيار مجلد كامل والبحث عن جميع الصور فيه"""
        folder = QFileDialog.getExistingDirectory(
            self,
            'اختر مجلد يحتوي على الصور'
        )
        
        if folder:
            # اختر/أنشئ مجلد السنة أولاً
            year_folder = choose_year_folder(self)
            if not year_folder:
                QMessageBox.warning(self, 'تنبيه', 'يجب اختيار أو إنشاء مجلد سنة')
                return

            # البحث عن جميع الصور في المجلد والمجلدات الفرعية
            from pathlib import Path
            import shutil, os
            folder_path = Path(folder)
            image_extensions = {'.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif', '.webp'}
            files = []

            # البحث في المجلد الحالي والمجلدات الفرعية
            for ext in image_extensions:
                files.extend([str(f) for f in folder_path.glob(f'*{ext}')])
                files.extend([str(f) for f in folder_path.glob(f'*{ext.upper()}')])
                files.extend([str(f) for f in folder_path.glob(f'**/*{ext}')])
                files.extend([str(f) for f in folder_path.glob(f'**/*{ext.upper()}')])

            if files:
                # إزالة التكرارات والترتيب
                found_files = sorted(list(set(files)))

                # نسخ الملفات إلى مجلد السنة
                dest_files = []
                for fp in found_files:
                    try:
                        basename = os.path.basename(fp)
                        dest = os.path.join(year_folder, basename)
                        shutil.copy2(fp, dest)
                        dest_files.append(dest)
                    except Exception:
                        pass

                self.selected_files = dest_files

                # إظهار عدد الصور المكتشفة
                count = len(self.selected_files)
                QMessageBox.information(
                    self,
                    'تم الاستيراد',
                    f'تم استيراد {count} صورة إلى مجلد السنة: {year_folder}'
                )

                self.update_list()
            else:
                QMessageBox.warning(self, 'تنبيه', 'لم يتم العثور على صور في المجلد')
    
    def update_list(self):
        """تحديث قائمة الملفات وتحليلها"""
        self.file_list.clear()
        info_text = 'تحليل الملفات:\n' + '='*50 + '\n'
        
        grouped = ImageSequenceHandler.group_images(
            [os.path.basename(f) for f in self.selected_files]
        )
        
        for filename in self.selected_files:
            basename = os.path.basename(filename)
            item = QListWidgetItem(basename)
            self.file_list.addItem(item)
            
            # تحليل الملف
            parsed = FilenameParser.parse_filename(basename)
            if parsed['is_valid']:
                info_text += f"\n📄 {basename}\n"
                info_text += f"  • الرقم: {parsed['number']}\n"
                info_text += f"  • التاريخ: {parsed['date']}\n"
                info_text += f"  • الجهة: {parsed['department']}\n"
                if parsed['sequence']:
                    info_text += f"  • التسلسل: {parsed['sequence']}\n"
        
        self.info_text.setText(info_text)
    
    def select_all_files(self):
        """تحديد جميع الملفات في القائمة"""
        for i in range(self.file_list.count()):
            self.file_list.item(i).setSelected(True)
        
        QMessageBox.information(self, 'تحديد', f'تم تحديد جميع الملفات ({self.file_list.count()} ملف)')
    
    def delete_selected(self):
        """حذف الملفات المحددة بعد التنبيه"""
        selected_items = self.file_list.selectedItems()
        
        if not selected_items:
            QMessageBox.warning(self, 'تنبيه', 'يجب تحديد ملفات للحذف أولاً')
            return
        
        count = len(selected_items)
        reply = QMessageBox.question(
            self,
            'تأكيد الحذف',
            f'هل تريد حذف {count} ملف من القائمة؟\n\nلن يتم حذفها من الحاسب، فقط من قائمة الاستيراد',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # احصل على أسماء الملفات المحددة
            selected_names = [item.text() for item in selected_items]
            
            # احذفها من selected_files
            self.selected_files = [
                f for f in self.selected_files 
                if os.path.basename(f) not in selected_names
            ]
            
            # احذفها من القائمة المرئية (احذف من الأسفل للأعلى لتجنب مشاكل الفهرسة)
            for i in range(self.file_list.count() - 1, -1, -1):
                if self.file_list.item(i) in selected_items:
                    self.file_list.takeItem(i)
            
            # احدّث المعلومات
            self.update_list()
            QMessageBox.information(self, 'نجح', f'تم حذف {count} ملف من القائمة')
    
    def get_files(self):
        """الحصول على الملفات المختارة"""
        return self.selected_files


class DestructionFormDialog(QDialog):
    """نافذة استمارة إتلاف الوثائق"""
    
    ROWS_PER_PAGE = 25  # عدد الصفوف في كل صفحة
    
    def __init__(self, parent=None, db=None, selected_docs=None):
        super().__init__(parent)
        self.db = db
        self.selected_docs = selected_docs or []
        self.setWindowTitle('استمارة إتلاف الوثائق')
        self.setMinimumSize(900, 700)
        self.init_ui()
        self.load_selected_documents()
        self.update_pages_info()
    
    def init_ui(self):
        """إنشاء واجهة المستخدم"""
        layout = QVBoxLayout()
        
        # عنوان النموذج
        title = QLabel('📋 استمارة إتلاف الوثائق')
        title.setStyleSheet('font-size: 18px; font-weight: bold; color: #2c3e50; padding: 10px;')
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        # معلومات الرأس
        header_group = QGroupBox('معلومات الجهة')
        header_layout = QFormLayout()
        
        self.agency_input = QLineEdit()
        self.agency_input.setPlaceholderText('أدخل اسم الوكالة')
        header_layout.addRow('الوكالة:', self.agency_input)
        
        self.directorate_input = QLineEdit()
        self.directorate_input.setPlaceholderText('أدخل اسم التشكيل أو المديرية')
        header_layout.addRow('التشكيل/المديرية:', self.directorate_input)
        
        self.section_input = QLineEdit()
        self.section_input.setPlaceholderText('أدخل اسم القسم')
        header_layout.addRow('القسم:', self.section_input)
        
        self.division_input = QLineEdit()
        self.division_input.setPlaceholderText('أدخل اسم الشعبة')
        header_layout.addRow('الشعبة:', self.division_input)
        
        header_group.setLayout(header_layout)
        layout.addWidget(header_group)
        
        # معلومات الصفحات
        pages_info_layout = QHBoxLayout()
        self.pages_info_label = QLabel('')
        self.pages_info_label.setStyleSheet('font-size: 12px; color: #666; padding: 5px;')
        pages_info_layout.addWidget(self.pages_info_label)
        pages_info_layout.addStretch()
        layout.addLayout(pages_info_layout)
        
        # جدول الوثائق
        docs_group = QGroupBox('الوثائق المراد إتلافها')
        docs_layout = QVBoxLayout()
        
        # أزرار التحكم بالجدول
        table_buttons = QHBoxLayout()
        
        add_row_btn = QPushButton('➕ إضافة صف')
        add_row_btn.clicked.connect(self.add_row)
        table_buttons.addWidget(add_row_btn)
        
        remove_row_btn = QPushButton('➖ حذف صف')
        remove_row_btn.clicked.connect(self.remove_row)
        table_buttons.addWidget(remove_row_btn)
        
        table_buttons.addStretch()
        
        load_selected_btn = QPushButton('📥 تحميل الوثائق المحددة')
        load_selected_btn.clicked.connect(self.load_selected_documents)
        table_buttons.addWidget(load_selected_btn)
        
        docs_layout.addLayout(table_buttons)
        
        # الجدول
        self.docs_table = QTableWidget()
        self.docs_table.setColumnCount(7)
        self.docs_table.setHorizontalHeaderLabels([
            'ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف الوثيقة\n(أ، ب، ج)', 'الفقرة القانونية'
        ])
        self.docs_table.setColumnWidth(0, 40)
        self.docs_table.setColumnWidth(1, 100)
        self.docs_table.setColumnWidth(2, 100)
        self.docs_table.setColumnWidth(3, 120)
        self.docs_table.setColumnWidth(4, 200)
        self.docs_table.setColumnWidth(5, 80)
        self.docs_table.setColumnWidth(6, 150)
        self.docs_table.setAlternatingRowColors(True)
        self.docs_table.model().rowsInserted.connect(self.update_pages_info)
        self.docs_table.model().rowsRemoved.connect(self.update_pages_info)
        
        docs_layout.addWidget(self.docs_table)
        docs_group.setLayout(docs_layout)
        layout.addWidget(docs_group)
        
        # أزرار الإجراءات
        buttons_layout = QHBoxLayout()
        
        print_btn = QPushButton('🖨️ طباعة')
        print_btn.setStyleSheet('background-color: #9b59b6; color: white; padding: 10px; font-size: 14px;')
        print_btn.clicked.connect(self.print_form)
        buttons_layout.addWidget(print_btn)
        
        export_excel_btn = QPushButton('📊 تصدير Excel')
        export_excel_btn.setStyleSheet('background-color: #27ae60; color: white; padding: 10px; font-size: 14px;')
        export_excel_btn.clicked.connect(self.export_to_excel)
        buttons_layout.addWidget(export_excel_btn)
        
        export_word_btn = QPushButton('📄 تصدير Word')
        export_word_btn.setStyleSheet('background-color: #3498db; color: white; padding: 10px; font-size: 14px;')
        export_word_btn.clicked.connect(self.export_to_word)
        buttons_layout.addWidget(export_word_btn)
        
        buttons_layout.addStretch()
        
        close_btn = QPushButton('❌ إغلاق')
        close_btn.clicked.connect(self.reject)
        buttons_layout.addWidget(close_btn)
        
        layout.addLayout(buttons_layout)
        self.setLayout(layout)
    
    def update_pages_info(self):
        """تحديث معلومات الصفحات"""
        total_docs = self.docs_table.rowCount()
        total_pages = (total_docs + self.ROWS_PER_PAGE - 1) // self.ROWS_PER_PAGE if total_docs > 0 else 1
        self.pages_info_label.setText(f'📄 إجمالي الوثائق: {total_docs} | عدد الصفحات: {total_pages} (25 وثيقة لكل صفحة)')
    
    def add_row(self):
        """إضافة صف جديد"""
        row = self.docs_table.rowCount()
        self.docs_table.insertRow(row)
        self.docs_table.setItem(row, 0, QTableWidgetItem(str(row + 1)))
    
    def remove_row(self):
        """حذف الصف المحدد"""
        current_row = self.docs_table.currentRow()
        if current_row >= 0:
            self.docs_table.removeRow(current_row)
            # تحديث أرقام التسلسل
            for i in range(self.docs_table.rowCount()):
                self.docs_table.setItem(i, 0, QTableWidgetItem(str(i + 1)))
    
    def load_selected_documents(self):
        """تحميل الوثائق المحددة من النافذة الرئيسية"""
        if not self.selected_docs:
            return
        
        self.docs_table.setRowCount(0)
        
        for idx, doc in enumerate(self.selected_docs):
            row = self.docs_table.rowCount()
            self.docs_table.insertRow(row)
            
            # ت (التسلسل)
            self.docs_table.setItem(row, 0, QTableWidgetItem(str(idx + 1)))
            
            # رقم الوثيقة
            doc_name = doc[1] or ''
            doc_number = doc_name.split()[0] if doc_name else ''
            self.docs_table.setItem(row, 1, QTableWidgetItem(doc_number))
            
            # تاريخها
            self.docs_table.setItem(row, 2, QTableWidgetItem(doc[2] or ''))
            
            # جهة الإصدار
            self.docs_table.setItem(row, 3, QTableWidgetItem(doc[4] or ''))
            
            # مضمونها
            self.docs_table.setItem(row, 4, QTableWidgetItem(doc[3] or ''))
            
            # تصنيف الوثيقة
            self.docs_table.setItem(row, 5, QTableWidgetItem(doc[5] or ''))
            
            # الفقرة القانونية
            self.docs_table.setItem(row, 6, QTableWidgetItem(doc[6] or ''))
        
        self.update_pages_info()
    
    def get_table_data(self):
        """جمع بيانات الجدول"""
        data = []
        for row in range(self.docs_table.rowCount()):
            row_data = []
            for col in range(self.docs_table.columnCount()):
                item = self.docs_table.item(row, col)
                row_data.append(item.text() if item else '')
            data.append(row_data)
        return data
    
    def split_data_into_pages(self, data):
        """تقسيم البيانات إلى صفحات (25 صف لكل صفحة)"""
        pages = []
        for i in range(0, len(data), self.ROWS_PER_PAGE):
            page_data = data[i:i + self.ROWS_PER_PAGE]
            # إضافة صفوف فارغة إذا كانت الصفحة غير مكتملة
            while len(page_data) < self.ROWS_PER_PAGE:
                empty_row = [str(len(page_data) + 1 + i)] + [''] * 6
                page_data.append(empty_row)
            pages.append(page_data)
        
        # إذا لم تكن هناك بيانات، أنشئ صفحة فارغة
        if not pages:
            page_data = []
            for j in range(self.ROWS_PER_PAGE):
                page_data.append([str(j + 1)] + [''] * 6)
            pages.append(page_data)
        
        return pages
    
    def generate_html_page(self, page_data, page_num, total_pages):
        """إنشاء HTML لصفحة واحدة - مصممة لتناسب ورقة A4 واحدة"""
        html = f'''
        <div style="page-break-after: always; page-break-inside: avoid; direction: rtl; font-family: Arial, sans-serif; width: 100%; max-height: 270mm;">
            <h3 style="text-align: center; margin: 0 0 5px 0; font-size: 14px; font-weight: bold;">استمارة إتلاف الوثائق</h3>
            
            <table style="width: 100%; margin-bottom: 3px; font-size: 9px; border: none;">
                <tr>
                    <td style="padding: 1px;">الوكالة: {self.agency_input.text()}</td>
                    <td style="padding: 1px;">التشكيل أو المديرية: {self.directorate_input.text()}</td>
                </tr>
                <tr>
                    <td style="padding: 1px;">القسم: {self.section_input.text()}</td>
                    <td style="padding: 1px;">الشعبة: {self.division_input.text()}</td>
                </tr>
            </table>
            
            <table style="width: 100%; border-collapse: collapse; font-size: 7px; table-layout: fixed;">
                <tr style="background-color: #D9E1F2;">
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 3%;">ت</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 8%;">رقم الوثيقة</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 10%;">تاريخها</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 12%;">جهة الإصدار</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 35%;">مضمونها</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 7%;">تصنيف</th>
                    <th style="border: 1px solid #000; padding: 2px; text-align: center; width: 25%;">الفقرة القانونية</th>
                </tr>
        '''
        
        for row in page_data:
            html += '<tr style="height: 8mm;">'
            for cell in row:
                html += f'<td style="border: 1px solid #000; padding: 1px; text-align: center; font-size: 7px;">{cell}</td>'
            html += '</tr>'
        
        html += f'''
            </table>
            <div style="text-align: left; font-size: 7px; margin-top: 2px;">صفحة {page_num} من {total_pages}</div>
        </div>
        '''
        return html
    
    def print_form(self):
        """طباعة الاستمارة مع معاينة باستخدام QPainter"""
        try:
            from PyQt6.QtPrintSupport import QPrinter, QPrintPreviewDialog
            from PyQt6.QtGui import QPageSize, QPageLayout
            from PyQt6.QtCore import QMarginsF
            
            # جمع البيانات وتقسيمها
            table_data = self.get_table_data()
            self._print_pages = self.split_data_into_pages(table_data)
            self._total_pages = len(self._print_pages)
            
            # حفظ معلومات النموذج
            self._form_info = {
                'agency': self.agency_input.text(),
                'directorate': self.directorate_input.text(),
                'section': self.section_input.text(),
                'division': self.division_input.text()
            }
            
            # إعداد الطابعة
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            page_size = QPageSize(QPageSize.PageSizeId.A4)
            page_layout = QPageLayout(page_size, QPageLayout.Orientation.Portrait, QMarginsF(15, 15, 15, 15))
            printer.setPageLayout(page_layout)
            
            # إنشاء نافذة المعاينة
            preview = QPrintPreviewDialog(printer, self)
            preview.setWindowTitle(f'معاينة الطباعة - {self._total_pages} صفحة')
            preview.setMinimumSize(800, 600)
            
            preview.paintRequested.connect(self._draw_pages_with_painter)
            preview.exec()
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء الطباعة:\n{str(e)}')
    
    def _draw_pages_with_painter(self, printer):
        """رسم الصفحات باستخدام QPainter للتحكم الكامل"""
        from PyQt6.QtGui import QPainter, QFont, QColor, QPen
        from PyQt6.QtPrintSupport import QPrinter
        from PyQt6.QtCore import Qt, QRectF
        
        painter = QPainter()
        painter.begin(printer)
        
        # استخدام وحدة Millimeter للحصول على أبعاد حقيقية
        page_rect = printer.pageRect(QPrinter.Unit.Millimeter)
        width_mm = page_rect.width()
        height_mm = page_rect.height()
        
        # تحويل من مليمتر إلى بكسل
        dpi = printer.resolution()
        px_per_mm = dpi / 25.4
        
        width = width_mm * px_per_mm
        height = height_mm * px_per_mm
        
        # هوامش بالمليمتر ثم تحويلها
        margin_mm = 10
        margin = margin_mm * px_per_mm
        content_width = width - (2 * margin)
        
        # ارتفاع الصف (حوالي 7mm للصف)
        row_height = 7 * px_per_mm
        
        for page_idx, page_data in enumerate(self._print_pages):
            if page_idx > 0:
                printer.newPage()
            
            y_pos = margin
            
            # ===== الشعار (أعلى وسط) =====
            import os
            logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "MOI.png")
            if os.path.exists(logo_path):
                from PyQt6.QtGui import QImage
                logo = QImage(logo_path)
                if not logo.isNull():
                    # حجم الشعار المناسب
                    logo_size = 25 * px_per_mm  # حوالي 25mm - أكبر وأجمل
                    logo_x = margin + (content_width - logo_size) / 2  # وسط الصفحة
                    logo_y = y_pos
                    painter.drawImage(QRectF(logo_x, logo_y, logo_size, logo_size), logo)
            
            # ===== معلومات الوزارة (أعلى يمين) =====
            ministry_font = QFont("Arial")
            ministry_font.setPointSize(10)
            ministry_font.setBold(True)
            painter.setFont(ministry_font)
            painter.setPen(Qt.GlobalColor.black)
            
            ministry_lines = [
                "وزارة الداخلية",
                "وكالة الوزارة لشؤون الإدارية والمالية",
                "مديرية إدارة الموارد البشرية",
                "مديرية السجلات والوثائق"
            ]
            
            line_h = 5 * px_per_mm
            temp_y = y_pos
            for line in ministry_lines:
                ministry_rect = QRectF(margin + content_width * 0.55, temp_y, content_width * 0.45, line_h)
                painter.drawText(ministry_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, line)
                temp_y += line_h
            
            # ===== معلومات الإصدار (أعلى يسار) =====
            version_font = QFont("Arial")
            version_font.setPointSize(8)
            painter.setFont(version_font)
            
            # رسم كل سطر بشكل منفصل (التسمية ثم / ثم القيمة)
            version_data = [
                ("رقم الإصدار", "0.1"),
                ("سنة الإصدار", "2023"),
                ("رقم الترميز", "م.ب-س"),
                ("نموذج", "(37)")
            ]
            
            inner_y = y_pos
            for label, value in version_data:
                # رسم القيمة أولاً (على اليسار)
                val_rect = QRectF(margin, inner_y, 15 * px_per_mm, line_h)
                painter.drawText(val_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, value)
                # رسم /
                slash_rect = QRectF(margin + 15 * px_per_mm, inner_y, 3 * px_per_mm, line_h)
                painter.drawText(slash_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, "/")
                # رسم التسمية (على اليمين)
                lbl_rect = QRectF(margin + 18 * px_per_mm, inner_y, 25 * px_per_mm, line_h)
                painter.drawText(lbl_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, label)
                inner_y += line_h
            
            y_pos = temp_y + 5 * px_per_mm
            
            # ===== العنوان الرئيسي (وسط) =====
            title_font = QFont("Arial")
            title_font.setPointSize(16)
            title_font.setBold(True)
            painter.setFont(title_font)
            painter.setPen(Qt.GlobalColor.black)
            
            title_height = 10 * px_per_mm
            title_rect = QRectF(margin, y_pos, content_width, title_height)
            painter.drawText(title_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, "استمارة إتلاف الوثائق")
            y_pos += title_height + (4 * px_per_mm)
            
            # ===== معلومات النموذج =====
            info_font = QFont("Arial")
            info_font.setPointSize(10)
            painter.setFont(info_font)
            
            # رسم كل حقل بشكل منفصل (التسمية ثم : ثم القيمة)
            info_data = [
                ("الوكالة", self._form_info['agency']),
                ("التشكيل أو المديرية", self._form_info['directorate']),
                ("القسم", self._form_info['section']),
                ("الشعبة", self._form_info['division'])
            ]
            
            line_height = 5 * px_per_mm
            label_w = 32 * px_per_mm
            
            for label, value in info_data:
                # رسم التسمية (على اليمين)
                lbl_rect = QRectF(margin + content_width - label_w, y_pos, label_w, line_height)
                painter.drawText(lbl_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, label)
                # رسم :
                colon_rect = QRectF(margin + content_width - label_w - (4 * px_per_mm), y_pos, 4 * px_per_mm, line_height)
                painter.drawText(colon_rect, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter, ":")
                # رسم القيمة
                val_rect = QRectF(margin, y_pos, content_width - label_w - (6 * px_per_mm), line_height)
                painter.drawText(val_rect, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter, value)
                y_pos += line_height
            
            y_pos += 3 * px_per_mm
            
            # ===== الجدول =====
            col_widths = [0.04, 0.08, 0.10, 0.12, 0.34, 0.08, 0.24]
            headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف', 'الفقرة القانونية']
            
            # رسم رأس الجدول
            header_font = QFont("Arial")
            header_font.setPointSize(8)
            header_font.setBold(True)
            painter.setFont(header_font)
            
            # خلفية الرأس
            painter.setBrush(QColor(217, 225, 242))
            painter.setPen(QPen(Qt.GlobalColor.black, 1))
            painter.drawRect(QRectF(margin, y_pos, content_width, row_height))
            
            # رسم خلايا الرأس من اليمين لليسار
            x_pos = margin + content_width
            for header, col_w in zip(headers, col_widths):
                cell_width = content_width * col_w
                x_pos -= cell_width
                cell_rect = QRectF(x_pos, y_pos, cell_width, row_height)
                painter.drawRect(cell_rect)
                painter.drawText(cell_rect, Qt.AlignmentFlag.AlignCenter, header)
            
            y_pos += row_height
            
            # رسم صفوف البيانات
            data_font = QFont("Arial")
            data_font.setPointSize(7)
            painter.setFont(data_font)
            painter.setBrush(Qt.GlobalColor.white)
            
            for row_idx, row in enumerate(page_data, 1):
                x_pos = margin + content_width
                row_data = [str(row_idx)] + list(row[1:])
                
                for cell, col_w in zip(row_data, col_widths):
                    cell_width = content_width * col_w
                    x_pos -= cell_width
                    cell_rect = QRectF(x_pos, y_pos, cell_width, row_height)
                    painter.drawRect(cell_rect)
                    display_text = str(cell)[:45] if len(str(cell)) > 45 else str(cell)
                    painter.drawText(cell_rect.adjusted(2, 0, -2, 0), Qt.AlignmentFlag.AlignCenter, display_text)
                
                y_pos += row_height
            
            # ===== رقم الصفحة =====
            page_font = QFont("Arial")
            page_font.setPointSize(8)
            painter.setFont(page_font)
            page_num_text = f"صفحة {page_idx + 1} من {self._total_pages}"
            page_rect_bottom = QRectF(margin, height - margin - (5 * px_per_mm), content_width, 5 * px_per_mm)
            painter.drawText(page_rect_bottom, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, page_num_text)
        
        painter.end()
    
    def export_to_excel(self):
        """تصدير إلى ملف Excel (مع تقسيم الصفحات)"""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        except ImportError:
            QMessageBox.warning(self, 'خطأ', 'مكتبة openpyxl غير مثبتة!\nقم بتثبيتها: pip install openpyxl')
            return
        
        # اختيار مكان الحفظ
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'حفظ استمارة الإتلاف',
            'استمارة_اتلاف_الوثائق.xlsx',
            'Excel Files (*.xlsx)'
        )
        
        if not file_path:
            return
        
        try:
            # جمع البيانات وتقسيمها
            table_data = self.get_table_data()
            pages = self.split_data_into_pages(table_data)
            
            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # إزالة الورقة الافتراضية
            
            # إعداد الأنماط
            title_font = Font(size=16, bold=True)
            header_font = Font(size=11, bold=True)
            cell_font = Font(size=10)
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            right_align = Alignment(horizontal='right', vertical='center')
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
            
            for page_num, page_data in enumerate(pages, 1):
                ws = wb.create_sheet(title=f'صفحة {page_num}')
                ws.sheet_view.rightToLeft = True
                
                # العنوان الرئيسي
                ws.merge_cells('A1:G1')
                ws['A1'] = 'استمارة إتلاف الوثائق'
                ws['A1'].font = title_font
                ws['A1'].alignment = center_align
                ws.row_dimensions[1].height = 25
                
                # معلومات الرأس
                ws['A3'] = f'الوكالة: {self.agency_input.text()}'
                ws['A3'].alignment = right_align
                ws.merge_cells('A3:G3')
                
                ws['A4'] = f'التشكيل أو المديرية: {self.directorate_input.text()}'
                ws['A4'].alignment = right_align
                ws.merge_cells('A4:G4')
                
                ws['A5'] = f'القسم: {self.section_input.text()}'
                ws['A5'].alignment = right_align
                ws.merge_cells('A5:G5')
                
                ws['A6'] = f'الشعبة: {self.division_input.text()}'
                ws['A6'].alignment = right_align
                ws.merge_cells('A6:G6')
                
                # رقم الصفحة
                ws['G7'] = f'صفحة {page_num} من {len(pages)}'
                ws['G7'].alignment = Alignment(horizontal='left', vertical='center')
                
                # رؤوس الجدول
                headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف\n(أ،ب،ج)', 'الفقرة القانونية']
                header_row = table.rows[0]
                for idx, header in enumerate(headers):
                    cell = header_row.cells[idx]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                
                # عرض الأعمدة
                ws.column_dimensions['A'].width = 5
                ws.column_dimensions['B'].width = 12
                ws.column_dimensions['C'].width = 12
                ws.column_dimensions['D'].width = 18
                ws.column_dimensions['E'].width = 30
                ws.column_dimensions['F'].width = 10
                ws.column_dimensions['G'].width = 25
                
                # بيانات الجدول
                for row_idx, row_data in enumerate(page_data):
                    for col_idx, value in enumerate(row_data):
                        cell = ws.cell(row=row_idx + 1, column=col_idx, value=value)
                        cell.font = cell_font
                        cell.alignment = center_align
                        cell.border = thin_border
            
            wb.save(file_path)
            QMessageBox.information(self, 'نجح', f'تم تصدير {len(pages)} صفحة إلى:\n{file_path}')
            os.startfile(file_path)
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء التصدير:\n{str(e)}')
    
    def export_to_word(self):
        """تصدير إلى ملف Word (مع تقسيم الصفحات)"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            QMessageBox.warning(self, 'خطأ', 'مكتبة python-docx غير مثبتة!\nقم بتثبيتها: pip install python-docx')
            return
        
        # اختيار مكان الحفظ
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'حفظ استمارة الإتلاف',
            'استمارة_اتلاف_الوثائق.docx',
            'Word Files (*.docx)'
        )
        
        if not file_path:
            return
        
        try:
            # جمع البيانات وتقسيمها
            table_data = self.get_table_data()
            pages = self.split_data_into_pages(table_data)
            
            doc = Document()
            
            for page_num, page_data in enumerate(pages, 1):
                # تعيين اتجاه المستند من اليمين لليسار
                if page_num == 1:
                    section = doc.sections[0]
                else:
                    section = doc.add_section()
                
                sectPr = section._sectPr
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                sectPr.append(bidi)
                
                # العنوان
                title = doc.add_paragraph('استمارة إتلاف الوثائق')
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].bold = True
                title.runs[0].font.size = Pt(16)
                
                # معلومات الرأس
                doc.add_paragraph(f'الوكالة: {self.agency_input.text()}')
                doc.add_paragraph(f'التشكيل أو المديرية: {self.directorate_input.text()}')
                doc.add_paragraph(f'القسم: {self.section_input.text()}')
                doc.add_paragraph(f'الشعبة: {self.division_input.text()}')
                
                # رقم الصفحة
                page_info = doc.add_paragraph(f'صفحة {page_num} من {len(pages)}')
                page_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # الجدول
                table = doc.add_table(rows=len(page_data) + 1, cols=7)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # رؤوس الجدول
                headers = ['ت', 'رقم الوثيقة', 'تاريخها', 'جهة الإصدار', 'مضمونها', 'تصنيف\n(أ،ب،ج)', 'الفقرة القانونية']
                header_row = table.rows[0]
                for idx, header in enumerate(headers):
                    cell = header_row.cells[idx]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                
                # بيانات الجدول
                for row_idx, row_data in enumerate(page_data):
                    row = table.rows[row_idx + 1]
                    for col_idx, value in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = value
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)
                
                # إضافة فاصل صفحة إذا لم تكن آخر صفحة
                if page_num < len(pages):
                    doc.add_page_break()
            
            doc.save(file_path)
            QMessageBox.information(self, 'نجح', f'تم تصدير {len(pages)} صفحة إلى:\n{file_path}')
            os.startfile(file_path)
            
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء التصدير:\n{str(e)}')


class MainWindow(QMainWindow):
    """النافذة الرئيسية للتطبيق"""
    
    def __init__(self):
        super().__init__()
        self.db = DatabaseManager('documents.db')
        self.image_manager = ImageManager('documents')
        self.setWindowTitle('برنامج أرشفة الكتب الرسمية')
        self.setGeometry(0, 0, 1200, 700)
        
        # تطبيق الأسلوب
        self.setStyleSheet(MAIN_STYLESHEET)
        
        self.init_ui()
        self.load_documents()
    
    def init_ui(self):
        """إنشاء واجهة المستخدم"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout()
        
        # شريط البحث
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel('البحث:'))
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText('ابحث حسب الرقم أو التاريخ أو الجهة...')
        self.search_input.textChanged.connect(self.search_documents)
        search_layout.addWidget(self.search_input)
        
        # حقل البحث
        search_field_label = QLabel('البحث في:')
        search_layout.addWidget(search_field_label)
        self.search_field = QComboBox()
        self.search_field.addItems(['اسم الوثيقة', 'التاريخ', 'الجهة', 'التصنيف'])
        self.search_field.currentTextChanged.connect(self.search_documents)
        search_layout.addWidget(self.search_field)
        
        main_layout.addLayout(search_layout)
        
        # شريط الأدوات
        toolbar_layout = QHBoxLayout()
        
        add_btn = QPushButton('➕ إضافة وثيقة')
        add_btn.clicked.connect(self.add_document)
        toolbar_layout.addWidget(add_btn)
        
        import_btn = QPushButton('📁 استيراد صور')
        import_btn.clicked.connect(self.import_images)
        toolbar_layout.addWidget(import_btn)
        
        view_btn = QPushButton('👁️ عرض الوثيقة')
        view_btn.clicked.connect(self.view_document)
        toolbar_layout.addWidget(view_btn)
        
        delete_btn = QPushButton('🗑️ حذف')
        delete_btn.clicked.connect(self.delete_document)
        toolbar_layout.addWidget(delete_btn)
        
        destruction_form_btn = QPushButton('📋 استمارة إتلاف')
        destruction_form_btn.clicked.connect(self.open_destruction_form)
        toolbar_layout.addWidget(destruction_form_btn)
        
        toolbar_layout.addStretch()
        
        select_all_btn = QPushButton('✓ تحديد الكل')
        select_all_btn.clicked.connect(self.select_all_documents)
        toolbar_layout.addWidget(select_all_btn)
        
        delete_selected_btn = QPushButton('🗑️ حذف المحددة')
        delete_selected_btn.clicked.connect(self.delete_selected_documents)
        toolbar_layout.addWidget(delete_selected_btn)
        
        refresh_btn = QPushButton('🔄 تحديث')
        refresh_btn.clicked.connect(self.load_documents)
        toolbar_layout.addWidget(refresh_btn)
        main_layout.addLayout(toolbar_layout)

        # محتوى رئيسي: قائمة السنوات على اليسار والجدول على اليمين
        content_layout = QHBoxLayout()

        # قائمة السنوات
        self.years_list = QListWidget()
        self.years_list.setMaximumWidth(220)
        self.years_list.itemClicked.connect(self.on_year_selected)
        content_layout.addWidget(self.years_list)

        # جدول الوثائق
        self.documents_table = QTableWidget()
        self.documents_table.setColumnCount(8)
        self.documents_table.setHorizontalHeaderLabels([
            '☑', 'رقم الوثيقة', 'التاريخ', 'المضمون', 'جهة الإصدار', 'التصنيف', 'المادة القانونية', 'عدد الصور'
        ])
        self.documents_table.setColumnWidth(0, 40)  # Checkbox column
        self.documents_table.setColumnWidth(1, 100)
        self.documents_table.setColumnWidth(2, 100)
        self.documents_table.setColumnWidth(3, 200)
        self.documents_table.setColumnWidth(4, 150)
        self.documents_table.setColumnWidth(5, 100)
        self.documents_table.setColumnWidth(6, 180)
        self.documents_table.setColumnWidth(7, 80)
        self.documents_table.setAlternatingRowColors(True)
        self.documents_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.documents_table.setSelectionMode(QTableWidget.SelectionMode.MultiSelection)
        
        # ضع الجدول داخل تخطيط عمودي (للسماح بعناصر إضافية إن لزم)
        right_layout = QVBoxLayout()
        right_layout.addWidget(self.documents_table)
        content_layout.addLayout(right_layout)

        main_layout.addLayout(content_layout)

        central_widget.setLayout(main_layout)
    
    def load_documents(self, year_filter=None):
        """تحميل قائمة الوثائق. إذا تم تمرير `year_filter` (مثل '2025')، يعرض الوثائق المرتبطة بتلك السنة فقط."""
        # تحديث قائمة السنوات
        self.refresh_years()

        self.documents_table.setRowCount(0)
        documents = self.db.get_all_documents()

        # إذا تم تحديد سنة، احصل على معرفات الوثائق التي تحتوي صورها في مجلد تلك السنة
        filter_ids = None
        if year_filter:
            filter_ids = set(self.db.get_document_ids_by_image_year(year_filter))
        
        # Disable updates for better performance
        self.documents_table.setUpdatesEnabled(False)
        
        for idx, doc in enumerate(documents):
            # إذا يوجد فلتر سنة وتوثيقة غير موجودة ضمن تلك السنة، تجاهلها
            if filter_ids is not None and doc[0] not in filter_ids:
                continue
            row = self.documents_table.rowCount()
            self.documents_table.insertRow(row)
            
            # Checkbox column
            checkbox = QCheckBox()
            checkbox.setStyleSheet('margin-left: 10px;')
            self.documents_table.setCellWidget(row, 0, checkbox)
            
            # رقم الوثيقة (من اسم الوثيقة)
            doc_name = doc[1] or ''
            # استخراج الرقم من اسم الوثيقة (مثل: "65 في 23-3-2025" -> "65")
            doc_number = doc_name.split()[0] if doc_name else ''
            item = QTableWidgetItem(doc_number)
            item.setData(Qt.ItemDataRole.UserRole, doc[0])  # احفظ معرف الوثيقة
            self.documents_table.setItem(row, 1, item)
            
            # التاريخ
            self.documents_table.setItem(row, 2, QTableWidgetItem(doc[2] or ''))
            
            # المضمون (العنوان)
            self.documents_table.setItem(row, 3, QTableWidgetItem(doc[3] or ''))
            
            # جهة الإصدار
            self.documents_table.setItem(row, 4, QTableWidgetItem(doc[4] or ''))
            
            # التصنيف
            self.documents_table.setItem(row, 5, QTableWidgetItem(doc[5] or ''))
            
            # المادة القانونية
            self.documents_table.setItem(row, 6, QTableWidgetItem(doc[6] or ''))
            
            # عدد الصور
            images = self.db.get_document_images(doc[0])
            self.documents_table.setItem(row, 7, QTableWidgetItem(str(len(images))))
            
            # Process events every 50 rows to keep UI responsive
            if idx % 50 == 0:
                QApplication.processEvents()
        
        # Re-enable updates
        self.documents_table.setUpdatesEnabled(True)

    def refresh_years(self):
        """تحديث قائمة السنوات من مجلد `documents/`"""
        from pathlib import Path
        import os
        docs_dir = Path(self.image_manager.storage_dir)
        self.years_list.clear()
        # إضافة خيار عرض الكل
        all_item = QListWidgetItem('الكل')
        self.years_list.addItem(all_item)
        if docs_dir.exists():
            for d in sorted(docs_dir.iterdir()):
                if d.is_dir() and d.name.isdigit():
                    item = QListWidgetItem(d.name)
                    self.years_list.addItem(item)

    def on_year_selected(self, item):
        text = item.text()
        if text == 'الكل':
            self.load_documents(None)
        else:
            self.load_documents(text)
    
    def add_document(self):
        """إضافة وثيقة جديدة"""
        dialog = AddDocumentDialog(self, self.db, self.image_manager)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            data = dialog.get_data()
            
            if not data['doc_name']:
                QMessageBox.warning(self, 'خطأ', 'يجب إدخال اسم الوثيقة')
                return
            
            doc_id = self.db.add_document(
                data['doc_name'],
                data['doc_date'],
                data['doc_title'],
                data['issuing_dept'],
                data['doc_classification'],
                data['legal_paragraph']
            )
            
            # حفظ جميع الصور الممسوحة مع الوثيقة الرئيسية
            scanned_images = data.get('scanned_images', [])
            # attachment_details_dict هو قاموس بمفاتيح 1, 2, 3, ...
            # حيث المفتاح 1 = المرفق الأول (الصورة الثانية، فهرس 1 في scanned_images)
            attachment_details_dict = data.get('attachment_details_dict', {})
            
            print(f"[DEBUG] add_document: عدد الصور = {len(scanned_images)}")
            print(f"[DEBUG] add_document: attachment_details_dict = {attachment_details_dict}")
            print(f"[DEBUG] add_document: مفاتيح القاموس = {list(attachment_details_dict.keys())}")
            
            if scanned_images:
                saved_count = 0
                # استخلاص مجلد السنة إن اختاره المستخدم في الحوار
                selected_year = None
                year_folder = data.get('selected_year_folder')
                try:
                    from pathlib import Path
                    selected_year = Path(year_folder).name if year_folder else None
                except Exception:
                    selected_year = None

                # enumerate(scanned_images, 0) -> idx يبدأ من 0
                # scanned_images[0] = الصورة الأولى = الوثيقة الرئيسية
                # scanned_images[1] = الصورة الثانية = المرفق الأول -> attachment_details_dict[1]
                # scanned_images[2] = الصورة الثالثة = المرفق الثاني -> attachment_details_dict[2]
                
                for idx, image_path in enumerate(scanned_images):
                    if os.path.exists(image_path):
                        try:
                            notes_text = None
                            
                            print(f"[DEBUG] معالجة الصورة idx={idx}")
                            
                            # الصورة الأولى (idx=0) هي الوثيقة الرئيسية - تستخدم بيانات الوثيقة الرئيسية
                            # الصورة الثانية (idx=1) هي المرفق الأول - بياناتها في attachment_details_dict[1]
                            # الصورة الثالثة (idx=2) هي المرفق الثاني - بياناتها في attachment_details_dict[2]
                            
                            if idx == 0:
                                # الوثيقة الرئيسية - استخدم بيانات data
                                print(f"[DEBUG] idx=0: الوثيقة الرئيسية")
                                merged_data = {
                                    'doc_name': data['doc_name'],
                                    'doc_date': data['doc_date'],
                                    'doc_title': data['doc_title'],
                                    'issuing_dept': data.get('issuing_dept', ''),
                                    'doc_classification': data.get('doc_classification', ''),
                                    'notes': ''
                                }
                            else:
                                # هذا مرفق - ابحث عن بياناته في القاموس
                                attachment_info = attachment_details_dict.get(idx, {})
                                print(f"[DEBUG] idx={idx}: المرفق {idx}, attachment_info = {attachment_info}")
                                
                                # تحقق هل هناك بيانات مخصصة
                                has_custom_data = False
                                if attachment_info and isinstance(attachment_info, dict):
                                    has_custom_data = any(
                                        v is not None and str(v).strip() != '' and v != 'اختر جهة الإصدار'
                                        for v in attachment_info.values()
                                    )
                                
                                print(f"[DEBUG] has_custom_data للمرفق {idx}: {has_custom_data}")
                                
                                if has_custom_data:
                                    # المرفق له بيانات مخصصة - استخدمها
                                    merged_data = {
                                        'doc_name': attachment_info.get('doc_name') or data['doc_name'],
                                        'doc_date': attachment_info.get('doc_date') or data['doc_date'],
                                        'doc_title': attachment_info.get('doc_title') or data['doc_title'],
                                        'issuing_dept': attachment_info.get('issuing_dept') or data.get('issuing_dept', ''),
                                        'doc_classification': attachment_info.get('doc_classification') or data.get('doc_classification', ''),
                                        'notes': attachment_info.get('notes', '')
                                    }
                                    print(f"[DEBUG] استخدام بيانات مخصصة للمرفق {idx}")
                                else:
                                    # المرفق ليس له بيانات مخصصة - استخدم بيانات الوثيقة الرئيسية
                                    merged_data = {
                                        'doc_name': data['doc_name'],
                                        'doc_date': data['doc_date'],
                                        'doc_title': data['doc_title'],
                                        'issuing_dept': data.get('issuing_dept', ''),
                                        'doc_classification': data.get('doc_classification', ''),
                                        'notes': ''
                                    }
                                    print(f"[DEBUG] استخدام بيانات الوثيقة الرئيسية للمرفق {idx}")
                            
                            print(f"[DEBUG] البيانات النهائية للصورة {idx}: {merged_data}")
                            
                            # إنشاء نص الملاحظات
                            notes_parts = []
                            if merged_data.get('doc_name'):
                                notes_parts.append(f"رقم: {merged_data['doc_name']}")
                            if merged_data.get('doc_date'):
                                notes_parts.append(f"تاريخ: {merged_data['doc_date']}")
                            if merged_data.get('doc_title'):
                                notes_parts.append(f"مضمون: {merged_data['doc_title']}")
                            if merged_data.get('issuing_dept'):
                                notes_parts.append(f"جهة: {merged_data['issuing_dept']}")
                            if merged_data.get('doc_classification'):
                                notes_parts.append(f"تصنيف: {merged_data['doc_classification']}")
                            if merged_data.get('notes'):
                                notes_parts.append(f"ملاحظات: {merged_data['notes']}")
                            
                            if notes_parts:
                                notes_text = " | ".join(notes_parts)
                            
                            # حفظ الصورة
                            saved_path = self.image_manager.save_image(
                                image_path,
                                doc_id,
                                idx + 1,  # page_number يبدأ من 1
                                year=selected_year
                            )
                            
                            print(f"[DEBUG] ✅ حفظ الصورة {idx} بـ notes: {notes_text}")
                            
                            # حفظ في قاعدة البيانات
                            self.db.add_image(
                                doc_id,
                                saved_path,
                                os.path.basename(image_path),
                                idx + 1,  # page_number يبدأ من 1
                                None,
                                1,
                                notes_text
                            )
                            
                            saved_count += 1
                                
                        except Exception as e:
                            print(f"خطأ في حفظ الصورة {idx}: {str(e)}")
                
                if saved_count > 0:
                    msg = f'تم الحفظ بنجاح!\n\n'
                    msg += f'الوثيقة الرئيسية مع {saved_count} صورة/مرفق'
                    QMessageBox.information(self, 'نجح ✅', msg)
                else:
                    QMessageBox.warning(self, 'تحذير', 'تم حفظ الوثيقة لكن لم يتم حفظ أي صورة')
            elif data.get('scanned_image') and os.path.exists(data['scanned_image']):
                # حفظ صورة واحدة فقط (للتوافق مع الكود القديم)
                try:
                    # استخلاص السنة ولوحة المستخدم
                    selected_year = None
                    year_folder = data.get('selected_year_folder')
                    try:
                        from pathlib import Path
                        selected_year = Path(year_folder).name if year_folder else None
                    except Exception:
                        selected_year = None

                    saved_path = self.image_manager.save_image(
                        data['scanned_image'],
                        doc_id,
                        1,
                        year=selected_year
                    )
                    
                    self.db.add_image(
                        doc_id,
                        saved_path,
                        os.path.basename(data['scanned_image']),
                        1,
                        None,
                        1,
                        None
                    )
                    
                    QMessageBox.information(self, 'نجح', 'تم إضافة الوثيقة والصورة بنجاح ✅')
                except Exception as e:
                    QMessageBox.warning(self, 'تحذير', f'تم حفظ الوثيقة لكن حدث خطأ في حفظ الصورة:\n{str(e)}')
            else:
                QMessageBox.information(self, 'نجح', 'تم إضافة الوثيقة بنجاح')
            
            self.load_documents()
    
    def import_images(self):
        """استيراد الصور"""
        dialog = ImportImagesDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            files = dialog.get_files()
            if not files:
                return
            
            # سؤال المستخدم عن استخراج المضمون من الصور (فقط إذا كان OCR متاحاً)
            extract_title = False
            ocr = None
            
            if OCR_AVAILABLE:
                # محاولة تهيئة OCR أولاً للتحقق من توفره
                try:
                    test_ocr = OCRExtractor()
                    if test_ocr.reader:
                        extract_title = QMessageBox.question(
                            self, 'استخراج المضمون',
                            '🔍 هل تريد استخراج المضمون (الموضوع) تلقائياً من الصور؟\n\n'
                            '• نعم: سيتم قراءة النص من الصور والبحث عن كلمة "الموضوع"\n'
                            '• لا: سيتم الاستيراد بدون استخراج المضمون\n\n'
                            '⚠️ ملاحظة: قد يستغرق الاستخراج وقتاً أطول',
                            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                        ) == QMessageBox.StandardButton.Yes
                        
                        if extract_title:
                            ocr = test_ocr
                    else:
                        # Tesseract غير مثبت - عرض رسالة للمستخدم
                        QMessageBox.information(
                            self, 'ميزة استخراج المضمون',
                            '📝 لتفعيل ميزة استخراج المضمون تلقائياً:\n\n'
                            '1. حمّل Tesseract OCR من:\n'
                            '   https://github.com/UB-Mannheim/tesseract/wiki\n\n'
                            '2. ثبته واختر اللغة العربية أثناء التثبيت\n\n'
                            '3. أعد تشغيل البرنامج\n\n'
                            'سيتم الاستيراد بدون استخراج المضمون حالياً.'
                        )
                except Exception as e:
                    print(f"[OCR] خطأ في التحقق من OCR: {str(e)}")
            
            # تحليل الملفات واستخراج البيانات
            documents_to_add = {}
            unrecognized = []
            
            for file_path in files:
                filename = os.path.basename(file_path)
                parsed = FilenameParser.parse_filename(filename)
                
                # تحقق إذا كان حرف "ص" أو "و" في اسم الملف
                default_dept = None
                if 'ص' in filename:
                    default_dept = 'شعبة أمن الأفراد عنة'
                elif 'و' in filename:
                    default_dept = 'قسم أمن الأفراد الأنبار'
                
                if parsed['is_valid']:
                    # استخدم الرقم والتاريخ والجهة كمعرّف فريد
                    doc_key = f"{parsed['number']}_{parsed['date']}_{parsed['department']}"
                    
                    # إذا كانت هذه أول مرة نرى هذا المفتاح، أنشئ وثيقة جديدة
                    if doc_key not in documents_to_add:
                        documents_to_add[doc_key] = {
                            'data': {
                                'doc_name': f"{parsed['number']} في {parsed['date']}",
                                'doc_date': parsed['date'],
                                'doc_title': '',
                                'issuing_dept': parsed['department'],
                                'doc_classification': '',
                                'legal_paragraph': ''
                            },
                            'images': []
                        }
                    
                    # أضف الصورة للوثيقة
                    documents_to_add[doc_key]['images'].append({
                        'path': file_path,
                        'filename': filename,
                        'sequence': parsed.get('sequence')
                    })
                else:
                    # الملفات التي لم يتم التعرف على صيغتها
                    # لكن إذا كانت تحتوي على ص أو و، أضفها بجهة الإصدار الافتراضية
                    if default_dept:
                        unrecognized.append((filename, default_dept))
                    else:
                        unrecognized.append((filename, None))
            
            # إذا كانت هناك ملفات غير معترف بها، اسأل المستخدم
            if unrecognized:
                only_with_dept = [f for f, d in unrecognized if d is not None]
                without_dept = [f for f, d in unrecognized if d is None]
                
                msg = f"عدد الملفات غير المعترف بصيغتها: {len(unrecognized)}\n\n"
                
                if only_with_dept:
                    msg += f"ملفات ستُضاف بجهة الإصدار: {only_with_dept[0].split()[0] if 'ص' in only_with_dept[0] else 'قسم أمن الأفراد الأنبار'} ({len(only_with_dept)})\n"
                
                if without_dept:
                    msg += f"ملفات بدون جهة إصدار ({len(without_dept)})\n\n"
                    msg += "هل تريد إضافة جميع الملفات؟\n\n"
                    msg += "أول 10 ملفات:\n"
                    for f, _ in unrecognized[:10]:
                        msg += f"• {f}\n"
                
                reply = QMessageBox.question(self, 'ملفات غير معترف بها', msg)
                
                if reply == QMessageBox.StandardButton.Yes:
                    # أنشئ وثائق للملفات غير المعترف بها
                    
                    # 1. ملفات بجهة الإصدار المحددة
                    dept_groups = {}
                    for filename, dept in unrecognized:
                        if dept:
                            if dept not in dept_groups:
                                dept_groups[dept] = []
                            dept_groups[dept].append(filename)
                    
                    for dept, filenames in dept_groups.items():
                        doc_key = f"unrecognized_{dept}"
                        documents_to_add[doc_key] = {
                            'data': {
                                'doc_name': f'ملفات مستورة عن {dept}',
                                'doc_date': '',
                                'doc_title': '',
                                'issuing_dept': dept,
                                'doc_classification': '',
                                'legal_paragraph': ''
                            },
                            'images': []
                        }
                        
                        for filename in filenames:
                            for file_path in files:
                                if os.path.basename(file_path) == filename:
                                    documents_to_add[doc_key]['images'].append({
                                        'path': file_path,
                                        'filename': filename,
                                        'sequence': None
                                    })
                                    break
                    
                    # 2. ملفات بدون جهة إصدار
                    no_dept_files = [f for f, d in unrecognized if d is None]
                    if no_dept_files:
                        doc_key = 'unrecognized_files_no_dept'
                        documents_to_add[doc_key] = {
                            'data': {
                                'doc_name': 'ملفات مستورة (بدون معلومات)',
                                'doc_date': '',
                                'doc_title': '',
                                'issuing_dept': '',
                                'doc_classification': '',
                                'legal_paragraph': ''
                            },
                            'images': []
                        }
                        
                        for filename in no_dept_files:
                            for file_path in files:
                                if os.path.basename(file_path) == filename:
                                    documents_to_add[doc_key]['images'].append({
                                        'path': file_path,
                                        'filename': filename,
                                        'sequence': None
                                    })
                                    break
            
            # حفظ الوثائق والصور في قاعدة البيانات
            # Calculate total images for progress
            total_images = sum(len(doc_info['images']) for doc_info in documents_to_add.values())
            
            # Create progress dialog
            progress_text = 'جاري استيراد الصور...'
            if extract_title:
                progress_text = 'جاري استيراد الصور واستخراج المضمون...'
            
            progress = QProgressDialog(progress_text, 'إلغاء', 0, total_images, self)
            progress.setWindowTitle('استيراد الصور')
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.show()
            
            imported_count = 0
            current_progress = 0
            extracted_titles_count = 0
            
            for doc_key, doc_info in documents_to_add.items():
                if progress.wasCanceled():
                    break
                    
                if not doc_info['images']:
                    continue
                
                # استخراج المضمون من الصور إذا طُلب ذلك (البحث في جميع الصور حتى نجد الموضوع)
                doc_title = doc_info['data']['doc_title']
                if extract_title and ocr and not doc_title and doc_info['images']:
                    progress.setLabelText(f'جاري استخراج المضمون من الصور...')
                    QApplication.processEvents()
                    
                    # البحث في جميع الصور حتى نجد الموضوع
                    for img_idx, img_info in enumerate(doc_info['images']):
                        try:
                            image_path = img_info['path']
                            print(f"[OCR] محاولة استخراج المضمون من الصورة {img_idx + 1}...")
                            
                            extracted_info = ocr.extract_document_info(image_path)
                            if extracted_info and extracted_info.get('doc_title'):
                                doc_title = extracted_info['doc_title']
                                doc_info['data']['doc_title'] = doc_title
                                extracted_titles_count += 1
                                print(f"[OCR] تم استخراج المضمون من الصورة {img_idx + 1}: {doc_title[:50]}...")
                                break  # توقف بعد إيجاد الموضوع
                        except Exception as e:
                            print(f"[OCR ERROR] خطأ في الصورة {img_idx + 1}: {str(e)}")
                            continue
                
                # استخراج رقم الوثيقة من الاسم (الجزء قبل كلمة "في")
                doc_name_parts = doc_info['data']['doc_name'].split(' في ')
                doc_number = doc_name_parts[0].strip() if doc_name_parts else ''
                doc_date = doc_info['data']['doc_date']
                
                print(f"[DEBUG] البحث عن وثيقة: رقم={doc_number}, تاريخ={doc_date}, اسم={doc_info['data']['doc_name']}")
                
                # تحقق من وجود الوثيقة بنفس الرقم والتاريخ
                existing = None
                if doc_number and doc_date:
                    existing = self.db.find_document_by_number_and_date(doc_number, doc_date)
                    print(f"[DEBUG] نتيجة البحث: {len(existing) if existing else 0} وثيقة")
                
                if existing:
                    doc_id = existing[0][0]
                    print(f"[DEBUG] تم إيجاد وثيقة موجودة: ID={doc_id}")
                    # تحديث المضمون إذا تم استخراجه
                    if doc_title:
                        self.db.update_document(doc_id, doc_title=doc_title)
                else:
                    # أنشئ وثيقة جديدة
                    print(f"[DEBUG] إنشاء وثيقة جديدة...")
                    doc_id = self.db.add_document(
                        doc_info['data']['doc_name'],
                        doc_info['data']['doc_date'],
                        doc_info['data']['doc_title'],
                        doc_info['data']['issuing_dept'],
                        doc_info['data']['doc_classification'],
                        doc_info['data']['legal_paragraph']
                    )
                    print(f"[DEBUG] تم إنشاء وثيقة جديدة: ID={doc_id}")
                
                # الحصول على عدد الصور الموجودة مسبقاً في الوثيقة
                existing_images = self.db.get_document_images(doc_id)
                start_img_idx = len(existing_images) + 1  # البدء من بعد آخر صورة
                
                # حفظ الصور
                for img_idx, img_info in enumerate(doc_info['images'], start_img_idx):
                    if progress.wasCanceled():
                        break
                    
                    current_progress += 1
                    progress.setValue(current_progress)
                    progress.setLabelText(f'جاري استيراد الصورة {current_progress} من {total_images}...')
                    QApplication.processEvents()  # Keep UI responsive
                    
                    try:
                        # حفظ الصورة في المجلد
                        saved_path = self.image_manager.save_image(
                            img_info['path'],
                            doc_id,
                            img_idx
                        )
                        
                        # أضف معلومات الصورة في قاعدة البيانات
                        self.db.add_image(
                            doc_id,
                            saved_path,
                            img_info['filename'],
                            img_idx,
                            img_info['sequence'],
                            1,
                            None
                        )
                        
                        imported_count += 1
                    
                    except Exception as e:
                        print(f"[ERROR] خطأ في حفظ الصورة {img_info['filename']}: {str(e)}")
            
            progress.setValue(total_images)
            progress.close()
            
            # الرسالة النهائية
            msg = f"✅ تم استيراد {imported_count} صورة بنجاح\n"
            msg += f"في {len(documents_to_add)} وثيقة"
            
            if extract_title and extracted_titles_count > 0:
                msg += f"\n\n📝 تم استخراج المضمون من {extracted_titles_count} وثيقة"
            
            if unrecognized:
                msg += f"\n\n⚠️ تم تخطي {len(unrecognized)} ملف"
            
            QMessageBox.information(self, 'نجح', msg)
            self.load_documents()
    
    def view_document(self):
        """عرض تفاصيل الوثيقة والصور"""
        current_row = self.documents_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, 'تنبيه', 'يجب اختيار وثيقة أولاً')
            return
        
        # احصل على معرف الوثيقة من البيانات المخزنة في الصف (column 1 now)
        doc_id_item = self.documents_table.item(current_row, 1)
        if not doc_id_item or not doc_id_item.data(Qt.ItemDataRole.UserRole):
            QMessageBox.warning(self, 'خطأ', 'لم يتم العثور على معرف الوثيقة')
            return
        
        doc_id = doc_id_item.data(Qt.ItemDataRole.UserRole)
        doc = self.db.get_document_by_id(doc_id)
        
        if doc:
            # الحصول على الصور
            images = self.db.get_document_images(doc_id)
            
            if not images:
                QMessageBox.warning(
                    self, 'تنبيه',
                    'لا توجد صور لهذه الوثيقة\n\nأرجو استيراد الصور أولاً'
                )
                return
            
            # استخراج مسارات الصور
            # جمع بيانات الصور مع معلومات المرفقات
            # هيكل جدول images: (0:id, 1:document_id, 2:image_path, 3:original_filename, 
            #                    4:page_number, 5:image_number, 6:sides, 7:created_date, 8:notes)
            images_data = []
            for img in images:
                img_path = img[2]  # العمود 2 هو image_path
                if os.path.exists(img_path):
                    notes_value = img[8] if len(img) > 8 else None  # العمود 8 هو notes
                    print(f"[DEBUG] img[8] (notes) = {notes_value}")
                    images_data.append({
                        'path': img_path,
                        'page_number': img[4] if len(img) > 4 else 0,
                        'notes': notes_value
                    })
            
            image_paths = [img['path'] for img in images_data]
            
            print(f"\n[MAIN] فتح عارض الوثائق:")
            print(f"  • معرف الوثيقة: {doc_id}")
            print(f"  • اسم الوثيقة: {doc[1]}")
            print(f"  • عدد الصور المسجلة: {len(images)}")
            print(f"  • عدد الصور الموجودة: {len(image_paths)}")
            for i, img_d in enumerate(images_data):
                print(f"  • صورة {i+1}: notes = {img_d.get('notes', 'لا يوجد')[:50] if img_d.get('notes') else 'لا يوجد'}...")
            if image_paths:
                print(f"  • أول صورة: {image_paths[0]}")
                print(f"  • آخر صورة: {image_paths[-1]}")
            
            if not image_paths:
                QMessageBox.warning(
                    self, 'خطأ',
                    'لا يمكن العثور على ملفات الصور\nقد تم حذفها من الحاسب'
                )
                return
            
            # فتح نافذة العرض مع بيانات الصور الكاملة
            try:
                viewer = DocumentViewerWindow(doc_id, doc, images_data, self)
                viewer.show()
                self.viewer_windows = getattr(self, 'viewer_windows', [])
                self.viewer_windows.append(viewer)
            except Exception as e:
                QMessageBox.critical(self, 'خطأ', f'خطأ في فتح العارض: {str(e)}')
    
    def delete_document(self):
        """حذف وثيقة"""
        current_row = self.documents_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, 'تنبيه', 'يجب اختيار وثيقة أولاً')
            return
        
        # احصل على معرف الوثيقة من UserRole (column 1 now)
        doc_id_item = self.documents_table.item(current_row, 1)
        if not doc_id_item or not doc_id_item.data(Qt.ItemDataRole.UserRole):
            QMessageBox.warning(self, 'خطأ', 'لم يتم العثور على معرف الوثيقة')
            return
        
        doc_id = doc_id_item.data(Qt.ItemDataRole.UserRole)
        
        reply = QMessageBox.question(
            self, 'تأكيد الحذف',
            'هل أنت متأكد من حذف هذه الوثيقة؟',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.db.delete_document(doc_id)
            QMessageBox.information(self, 'نجح', 'تم حذف الوثيقة')
            self.load_documents()
    
    def open_destruction_form(self):
        """فتح نافذة استمارة إتلاف الوثائق"""
        # الحصول على الوثائق المحددة
        selected_docs = []
        for row in range(self.documents_table.rowCount()):
            checkbox = self.documents_table.cellWidget(row, 0)
            if checkbox and checkbox.isChecked():
                doc_id_item = self.documents_table.item(row, 1)
                if doc_id_item:
                    doc_id = doc_id_item.data(Qt.ItemDataRole.UserRole)
                    if doc_id:
                        doc = self.db.get_document_by_id(doc_id)
                        if doc:
                            selected_docs.append(doc)
        
        # فتح النافذة
        dialog = DestructionFormDialog(self, self.db, selected_docs)
        dialog.exec()
    
    def select_all_documents(self):
        """تحديد جميع الوثائق"""
        for row in range(self.documents_table.rowCount()):
            checkbox = self.documents_table.cellWidget(row, 0)
            if checkbox:
                checkbox.setChecked(True)
    
    def delete_selected_documents(self):
        """حذف جميع الوثائق المحددة"""
        # Get checked rows
        checked_rows = []
        for row in range(self.documents_table.rowCount()):
            checkbox = self.documents_table.cellWidget(row, 0)
            if checkbox and checkbox.isChecked():
                checked_rows.append(row)
        
        if not checked_rows:
            QMessageBox.warning(self, 'تنبيه', 'يجب تحديد وثائق أولاً')
            return
        
        count = len(checked_rows)
        reply = QMessageBox.question(
            self,
            'تأكيد الحذف',
            f'هل أنت متأكد من حذف {count} وثيقة؟',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Get doc IDs from checked rows
            doc_ids = []
            for row in checked_rows:
                doc_id_item = self.documents_table.item(row, 1)  # Column 1 now has doc number
                if doc_id_item:
                    doc_id = doc_id_item.data(Qt.ItemDataRole.UserRole)
                    if doc_id:
                        doc_ids.append(doc_id)
            
            # Create progress dialog
            progress = QProgressDialog('جاري حذف الوثائق...', 'إلغاء', 0, len(doc_ids), self)
            progress.setWindowTitle('حذف الوثائق')
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.show()
            
            # Delete from database with progress
            deleted_count = 0
            for i, doc_id in enumerate(doc_ids):
                if progress.wasCanceled():
                    break
                progress.setValue(i)
                progress.setLabelText(f'جاري حذف الوثيقة {i + 1} من {len(doc_ids)}...')
                QApplication.processEvents()  # Keep UI responsive
                
                try:
                    self.db.delete_document(doc_id)
                    deleted_count += 1
                except Exception as e:
                    print(f'خطأ في حذف الوثيقة {doc_id}: {e}')
            
            progress.setValue(len(doc_ids))
            progress.close()
            
            # Reload table
            self.load_documents()
            QMessageBox.information(self, 'نجح', f'تم حذف {deleted_count} وثيقة بنجاح')
    
    def search_documents(self):
        """البحث عن الوثائق والمرفقات"""
        search_term = self.search_input.text().strip()
        if not search_term:
            self.load_documents()
            return
        
        # تحديد حقل البحث
        field_map = {
            'اسم الوثيقة': 'doc_name',
            'التاريخ': 'doc_date',
            'الجهة': 'issuing_dept',
            'التصنيف': 'doc_classification'
        }
        
        search_field = field_map.get(self.search_field.currentText(), 'doc_name')
        
        self.documents_table.setRowCount(0)
        
        # استخدام البحث الجديد الذي يشمل المرفقات
        results_dict = self.db.search_documents_and_attachments(search_term, search_field)
        
        # Disable updates for better performance
        self.documents_table.setUpdatesEnabled(False)
        
        for idx, (key, result_data) in enumerate(results_dict.items()):
            doc = result_data['doc']
            source = result_data['source']
            attachment_info = result_data['attachment_info']
            
            row = self.documents_table.rowCount()
            self.documents_table.insertRow(row)
            
            # Checkbox column
            checkbox = QCheckBox()
            checkbox.setStyleSheet('margin-left: 10px;')
            self.documents_table.setCellWidget(row, 0, checkbox)
            
            # رقم الوثيقة/المرفق
            if source == 'attachment' and attachment_info:
                # استخراج المعلومات من ملاحظات المرفق
                doc_number = result_data['doc_number']
                # إضافة علامة للمرفق
                display_number = f"📎 {doc_number}" if doc_number else ''
            else:
                doc_name = doc[1] or ''
                doc_number = doc_name.split()[0] if doc_name else ''
                display_number = doc_number
            
            item = QTableWidgetItem(display_number)
            item.setData(Qt.ItemDataRole.UserRole, doc[0])  # احفظ معرف الوثيقة
            self.documents_table.setItem(row, 1, item)
            
            # التاريخ
            if source == 'attachment' and attachment_info:
                # استخراج التاريخ من ملاحظات المرفق
                import re
                date_match = re.search(r'تاريخ:\s*([^\|]+)', attachment_info)
                date_val = date_match.group(1).strip() if date_match else (doc[2] or '')
            else:
                date_val = doc[2] or ''
            self.documents_table.setItem(row, 2, QTableWidgetItem(date_val))
            
            # المضمون
            if source == 'attachment' and attachment_info:
                title_match = re.search(r'مضمون:\s*([^\|]+)', attachment_info)
                title_val = title_match.group(1).strip() if title_match else (doc[3] or '')
            else:
                title_val = doc[3] or ''
            self.documents_table.setItem(row, 3, QTableWidgetItem(title_val))
            
            # الجهة
            if source == 'attachment' and attachment_info:
                dept_match = re.search(r'جهة:\s*([^\|]+)', attachment_info)
                dept_val = dept_match.group(1).strip() if dept_match else (doc[4] or '')
            else:
                dept_val = doc[4] or ''
            self.documents_table.setItem(row, 4, QTableWidgetItem(dept_val))
            
            # التصنيف
            if source == 'attachment' and attachment_info:
                class_match = re.search(r'تصنيف:\s*([^\|]+)', attachment_info)
                class_val = class_match.group(1).strip() if class_match else (doc[5] or '')
            else:
                class_val = doc[5] or ''
            self.documents_table.setItem(row, 5, QTableWidgetItem(class_val))
            
            # المادة القانونية
            self.documents_table.setItem(row, 6, QTableWidgetItem(doc[6] or ''))
            
            # عدد الصور
            images = self.db.get_document_images(doc[0])
            self.documents_table.setItem(row, 7, QTableWidgetItem(str(len(images))))
            
            # Process events every 50 rows
            if idx % 50 == 0:
                QApplication.processEvents()
        
        # Re-enable updates
        self.documents_table.setUpdatesEnabled(True)


def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == '__main__':
    main()
